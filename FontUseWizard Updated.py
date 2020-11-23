#Last update: 07/10/20

import os
import shutil
import csv
import pyxj
import re
import struct
from pyxj import dialog
from pathlib import Path
from fontio3 import fontedit
from fontio3.CFF import CFF
from fontio3.utilities.filewalker import FileWalker
import plistlib
from urllib.request import urlopen
import unicodedata
from libfont3 import fontEditorFromFont
import mmap
from docx import Document
from docx.shared import Inches
from collections import defaultdict
from datetime import date
from zipfile import ZipFile
from enum import Enum
import string

class Usecase(Enum):
	App_Android = 1
	App_IOS = 2
	Website = 3
	DigitalAds = 4
	PDF = 5

SCRIPT_VERSION="v1.03"

def fix_nulls(s):
    for line in s:
        yield line.replace('\0', '')

def remove_control_characters(s):
    return "".join(ch for ch in s if unicodedata.category(ch)[0]!="C")

def initFontData():
	default_value = ''
	font_dict = dict.fromkeys(['Name','Remarks','Copyright','Trademark','Family','Subfamily','LicenseDesc','LicenseURL','UniqueID','Misc1','Misc2','Version','Manufacturer','Designer',
		'DesignerURL','Vendor','VendorURL','VendorID','FontFileName','FileType','Path','WebPath','Foundry','Enforce','WhyEnforce','UseCase','Title','AppPlatform','AppDev','AppURL','VerifiedDomain','LicenseIndicator'],default_value)
	return font_dict

def makeFontFromCSV(row):
	font = initFontData()
	
	font['Enforce'] = row[0]
	font['Name'] = row[1]
	font['UseCase'] = row[2]
	font['Title'] = row[3]
	font['FontFileName'] = row[4]
	font['Copyright'] = row[5]
	font['Trademark'] = row[6]
	font['Family'] = row[7]
	font['Subfamily'] = row[8]
	font['LicenseDesc'] = row[9]
	font['LicenseURL'] = row[10]
	font['UniqueID'] = row[11]
	font['Misc1'] = row[12]
	font['Misc2'] = row[13]
	font['Version'] = row[14]
	font['Manufacturer'] = row[15]
	font['Designer'] = row[16]
	font['DesignerURL'] = row[17]
	font['Vendor'] = row[18]
	font['VendorURL'] = row[19]
	font['VendorID'] = row[20]
	font['FileType'] = row[21]
	font['Foundry'] = row[22]
	font['Path'] = row[23]
	font['WebPath'] = row[24]
	font['VerifiedDomain'] = row[25]
	font['LicenseIndicator'] = row[26]
	font['AppPlatform'] = row[27]
	font['AppDev'] = row[28]
	font['AppURL'] = row[29]
	font['WhyEnforce'] = row[30]
	font['Remarks'] = row[31]
	
	return font

#write results from general scan to a CSV file
def write_dir_scan(fonts_list, path):
	outpath = path + 'FontInfo_' + SCRIPT_VERSION + '.csv'
	output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))
	#Write CSV headers
	output.writerow(['Enforceable IP?','Font Name','Font File Name','Copyright','Trademark','Family', 'Subfamily','License Description', 'License Info URL', 'Unique ID', 'Misc1','Misc2','Version', 'Manufacturer', 'Designer', 'Designer URL','Vendor', 'Vendor URL', 'Vendor ID','File Type','Path','WhyEnforce','Remarks'])
	
	
	for font in fonts_list:
		remarks = ''
		
		if 'Remarks' in font.keys():
			for remark in font['Remarks']:
				remarks = remarks + remark + ' '
		
		if font['VendorID'] != '':
			font['VendorID'] = font['VendorID'].decode('utf-8')
		
		output.writerow([font['Enforce'], font['Name'], font['FontFileName'], font['Copyright'], font['Trademark'], font['Family'], font['Subfamily'], font['LicenseDesc'], font['LicenseURL'], font['UniqueID'], font['Misc1'],font['Misc2'], font['Version'], font['Manufacturer'], font['Designer'], font['DesignerURL'], font['Vendor'], font['VendorURL'], font['VendorID'], font['FileType'],font['Path'], font['WhyEnforce'], remarks])
	
	print('Done! Results saved to ' + outpath)

#similar to write_dir_scan but includes extra information pertaining to use case
def create_font_use_csv(fonts_list, path):
	output = csv.writer(open(path,'w',encoding='utf8', newline = ''))
	#Write CSV headers
	output.writerow(['Enforceable IP?','Font Name','Use Case', 'Website/App Title', 'Font File Name','Copyright','Trademark','Family', 'Subfamily','License Description', 'License Info URL', 'Unique ID', 'Misc1','Misc2','Version', 'Manufacturer', 'Designer', 'Designer URL','Vendor', 'Vendor URL', 'Vendor ID','File Type','Foundry (test)','Path','Web Path','Verified Domain', 'License Indicator','App Platform', 'App Developer', 'App URL','WhyEnforce','Remarks'])

	waybackURLs = ['https://web.archive.org/web/20200406145022/',
					'https://web.archive.org/web/20200101072516/',
					'https://web.archive.org/web/20190702000358/',
					'https://web.archive.org/web/20190101081056/',
					'https://web.archive.org/web/20180701184757/',
					'https://web.archive.org/web/20180104141938/',
					'https://web.archive.org/web/20170701023847/',
					'https://web.archive.org/web/20170101023847/',
					'https://web.archive.org/web/20160711135651/',
					'https://web.archive.org/web/20160112094559/',
					'https://web.archive.org/web/20150706032903/',
					'https://web.archive.org/web/20150101071917/',
					'https://web.archive.org/web/20140711180223/',
					'https://web.archive.org/web/20140103013856/',
					'https://web.archive.org/web/20130708001814/',
					'https://web.archive.org/web/20130101103045/',
					'https://web.archive.org/web/20120602145257/',
					'https://web.archive.org/web/20120104183204/',
					'https://web.archive.org/web/20110608173044/',
					'https://web.archive.org/web/20110101110846/',
					'https://web.archive.org/web/20100102003419/']

	waybackSet = set()
	for font in fonts_list:
		remarks = ''
		
		if 'Remarks' in font.keys():
			for remark in font['Remarks']:
				remarks = remarks + remark + ' '
		for key in ['Enforce', 'Name','UseCase','Title','FontFileName','Copyright','Trademark','Family','Subfamily','LicenseDesc','LicenseURL','UniqueID','Misc1','Misc2','Version','Manufacturer','Designer','DesignerURL','Vendor','VendorURL','VendorID', 'FileType','Foundry','Path','WebPath','VerifiedDomain','LicenseIndicator','AppPlatform','AppDev','AppURL', 'WhyEnforce']:
			if key not in font:
				font[key] = ""
		if font['VendorID'] != '':
			font['VendorID'] = font['VendorID'].decode('utf-8')

		if font['UseCase'] == 'Web Font' and font['Enforce'] != 'No':
			#for url in waybackURLs:
				#waybackCSV.writerow([url+font['Title']])
			url = remove_prefix(font['Title'], "https")
			url = remove_prefix(url, "www.")
			waybackSet.add("https://www." + url)

		if font['UseCase'] == 'PDF':
			font['Title'] = os.path.basename(os.path.dirname(font['Path'])) + '.pdf'

		output.writerow([font['Enforce'], font['Name'], font['UseCase'], font['Title'], font['FontFileName'], font['Copyright'], font['Trademark'], font['Family'], font['Subfamily'], font['LicenseDesc'], font['LicenseURL'], font['UniqueID'], font['Misc1'],font['Misc2'], font['Version'], font['Manufacturer'], font['Designer'], font['DesignerURL'], font['Vendor'], font['VendorURL'], font['VendorID'], font['FileType'],font['Foundry'],font['Path'],font['WebPath'], font['VerifiedDomain'],font['LicenseIndicator'],font['AppPlatform'],font['AppDev'],font['AppURL'],font['WhyEnforce'],remarks])

		if len(waybackSet)>0:
			waybackCSVFile = os.path.join(os.path.dirname(path), 'wayback.csv')
			waybackCSV = csv.writer(open(waybackCSVFile, 'w', encoding='utf8', newline=''))
			for url in waybackSet:
				for wayback in waybackURLs:
					waybackCSV.writerow([wayback + url])

def remove_prefix(str, prefix):
    if str.startswith(prefix):
        return str[len(prefix):]
    else:
        return str

def create_font_use_doc(csv_path, dirpath, pre_csv_path = '', redirectedURLs = None):
	#prompt for client name
	client = pyxj.askString('Please enter the name of the client.    ' + SCRIPT_VERSION, 'Client: ', cancelable=False)
	
	with open(csv_path,'r', encoding="utf8") as myFile:
		reader = csv.reader(fix_nulls(myFile))
		rows = list(reader)
	
	document = Document()
	#set up document font
	docFont = document.styles['Normal'].font
	docFont.name = 'Helvetica Now Display Light'
	
	heading1Font = document.styles['Heading 1'].font
	heading1Font.name = 'Helvetica Now Display Light'
	
	heading2Font = document.styles['Heading 2'].font
	heading2Font.name = 'Helvetica Now Display Light'
	
	heading3Font = document.styles['Heading 3'].font
	heading3Font.name = 'Helvetica Now Display Light'
	
	section = document.sections[0]
	header = section.header
	paragraph = header.paragraphs[0]
	paragraph.text = 'Monotype Imaging\t\t' + date.today().isoformat()
	paragraph.style = document.styles["Header"]
	
	iOS_apps = []
	android_apps = []
	websites = []
	metadata_fonts = []
	fonts_list = []
	oldestURLs = {}
	
	document.add_paragraph('[REPLACE WITH TABLE OF CONTENTS]')

	for row in rows[1:]:
		fontData = makeFontFromCSV(row)
		if fontData['Enforce'] != 'No':
			fonts_list.append(fontData)

	#oldest URL Logic
	if pre_csv_path != '':
		waybackRegex = 'web.archive.org\/web\/\d{14}\/'
		fonts_list = sorted(fonts_list, key=lambda s: re.sub(waybackRegex, "", s['Title']))
		temp_font_list = []
		prevURL = ''
		for font in fonts_list:
			actualURL = re.sub(waybackRegex, "", font['Title'])
			if prevURL != actualURL or font == fonts_list[-1]:
				if len(temp_font_list) > 0:
					currURL = temp_font_list[len(temp_font_list) -1]['Title']
					temp_fontnames = []
					old_dict = {}
					temp_font_list.reverse()
					for tempfont in temp_font_list:
						if currURL == tempfont['Title']:
							temp_fontnames.append(tempfont['Name'])
						else:
							break
					temp_font_list.reverse()
					for temp_fontname in temp_fontnames:
						for tempfont in temp_font_list:
							if tempfont['Name'] == temp_fontname:
								if redirectedURLs and tempfont['Title'] in redirectedURLs:
									old_dict[temp_fontname] = redirectedURLs[tempfont['Title']]
								else:
									old_dict[temp_fontname] = tempfont['Title']
								break
					if prevURL not in oldestURLs:
						oldestURLs[prevURL] = {}
					oldestURLs[prevURL]['CurrentFonts'] = old_dict
				temp_font_list.clear()
				prevURL = actualURL
			temp_font_list.append(font)
			if actualURL not in oldestURLs:
				oldestURLs[actualURL] = {}
				if redirectedURLs and font['Title'] in redirectedURLs:
					oldestURLs[actualURL]['OldestURL'] = redirectedURLs[font['Title']]
				else:
					oldestURLs[actualURL]['OldestURL'] = font['Title']
		fonts_list = []
		with open(pre_csv_path, 'r', encoding="utf8") as myFile:
			reader = csv.reader(fix_nulls(myFile))
			rows = list(reader)
		for row in rows[1:]:
			fontData = makeFontFromCSV(row)
			if fontData['Enforce'] != 'No':
				fonts_list.append(fontData)

	iOS_apps = getIOSApps(fonts_list)
	android_apps, amazon_apps = getAndroidApps(fonts_list)
	websites = getWebsites(fonts_list, 'Web Font')
	ads = getWebsites(fonts_list, 'Digital Ad')
	pdfs = getWebsites(fonts_list, 'PDF')
	
	if len(fonts_list) > 0:
		#Write Docs
		if len(iOS_apps) > 0:
			document.add_heading(client + ' Mobile Applications [iOS]', level=1)
			for app in iOS_apps:
				document.add_heading(app[0], level=2)
				p = document.add_paragraph('Developer: ', style = 'List Bullet')
				p.add_run(app[1])
				p = document.add_paragraph('Web URL: ', style = 'List Bullet')
				p.add_run(app[2])
				document.add_paragraph("Embedded Font Software", style = 'List Bullet')
				for font in app[3]:
					if 'AR/VR' in font['UseCase']:
						document.add_paragraph('AR/VR Font - ' + font['Name'], style='List Bullet 2')
					else:
						document.add_paragraph(font['Name'], style = 'List Bullet 2')
				document.add_paragraph('Screenshot of Font Software Embedded Within Mobile Application', style = 'List Bullet')
				for path in set(app[4]):
					document.add_paragraph(path)
				document.add_paragraph('App Version History', style = 'List Bullet')
				app_annie_url = 'https://www.appannie.com/apps/ios/app/' + app[2].split('/')[-1].replace('id','') + '/details/'
				document.add_paragraph(app_annie_url)
				document.add_paragraph('Internal Notes', style = 'List Bullet')
				#check levels of enforcement
				enf_MT = False
				enf_3p = False
				for font in app[3]:
					if font['Enforce'] == 'Yes':
						enf_MT = True
					elif font['Enforce'] != 'Yes':
						enf_3p = True
				if enf_MT:
					document.add_paragraph("Embedded Monotype Font Software", style = 'List Bullet 2')
					for font in app[3]:
						if font['Enforce'] == 'Yes':
							document.add_paragraph(font['Name'], style = 'List Bullet 3')
				if enf_3p:
					document.add_paragraph("Other Embedded Font Software", style = 'List Bullet 2')
					for font in app[3]:
						if font['Enforce'] != 'Yes':
							document.add_paragraph(font['Name'], style = 'List Bullet 3')

		if len(android_apps) > 0:
			document.add_heading(client + ' Mobile Applications [Android]', level=1)
			write_android_apps(android_apps, document)

		if len(amazon_apps) > 0:
			document.add_heading(client + ' Mobile Applications [Amazon]', level=1)
			write_android_apps(amazon_apps, document)
				
		if len(websites) > 0:
			write_websites(client, document, oldestURLs, dirpath, pre_csv_path, websites, Usecase.Website)

		if len(ads) > 0:
			write_websites(client, document, oldestURLs, dirpath, pre_csv_path, ads, Usecase.DigitalAds)

		if len(pdfs) > 0:
			write_websites(client, document, oldestURLs, dirpath, pre_csv_path, pdfs, Usecase.PDF)
			
		
		#write metadata
		document.add_heading('Font Software Metadata', level = 1)
		
		#fonts_check: [filename, name]
		fonts_check = []
		fonts_check.append(['check','check'])
		metadata_fonts = [] 
		
		for font in fonts_list:
			found = False
			for checkfont in fonts_check:
				if font['FontFileName'] == checkfont[0] and font['Name'] == checkfont[1]:
					found = True
			if not found:
				fonts_check.append([font['FontFileName'],font['Name']])
				metadata_fonts.append(font)
		
		#sorting metadata so that it's lumped together by font name
		sorted_metadata = sorted(metadata_fonts, key = lambda k : k['Name'])
		disp_names = []
		for font in sorted_metadata:
			if font['Name'] not in disp_names:
				document.add_heading(font['Name'], level = 2)
				disp_names.append(font['Name'])
				
			table = document.add_table(rows = 17, cols = 2)
			table.style = 'TableGrid'
			hdr_cells = table.columns[0].cells
			for cell in hdr_cells:
				cell.width = Inches(0.8)
			hdr_cells[0].text = 'name id' #filename
			hdr_cells[1].text = 'Full Name'
			hdr_cells[2].text = 'Family'
			hdr_cells[3].text = 'Subfamily'
			hdr_cells[4].text = 'Copyright'
			hdr_cells[5].text = 'Trademark'
			hdr_cells[6].text = 'License Description'
			hdr_cells[7].text = 'License Info URL'
			hdr_cells[8].text = 'Unique ID'
			hdr_cells[9].text = 'Version'
			hdr_cells[10].text = 'Manufacturer'
			hdr_cells[11].text = 'Designer'
			hdr_cells[12].text = 'Designer URL'
			hdr_cells[13].text = 'Vendor'
			hdr_cells[14].text = 'Vendor URL'
			hdr_cells[15].text = 'Vendor ID'
			hdr_cells[16].text = 'Misc'

			fnt_cells = table.columns[1].cells
			prepare_doc_table(fnt_cells, font)

			document.add_paragraph('')
	
		document.add_heading('Internal Notes', level = 1)
		document.add_paragraph('')
		document.add_heading('Company Structure', level = 2)
		document.add_paragraph('')
		document.add_heading('Reviewed Websites', level = 2)
		document.add_paragraph('')
		document.add_heading('Reviewed Apps', level = 2)
		document.add_paragraph('')
		#document.add_heading('License Checks', level = 2)
		#document.add_paragraph('')
		#document.add_heading('SFDC', level = 3)
		#document.add_paragraph('')
		#document.add_heading('DC', level = 3)
		#document.add_paragraph('')
		outfileName = client + " Font Use (Internal Version).docx"
		outpath = os.path.join(os.path.dirname(csv_path), outfileName)
		document.save(outpath)


def prepare_doc_table(fnt_cells, font, handleException = False):
	if not handleException:
		try:
			fnt_cells[0].text = font['FontFileName']
			fnt_cells[1].text = font['Name']
			fnt_cells[2].text = font['Family']
			fnt_cells[3].text = font['Subfamily']
			fnt_cells[4].text = font['Copyright']
			fnt_cells[5].text = font['Trademark']
			fnt_cells[6].text = font['LicenseDesc']
			fnt_cells[7].text = font['LicenseURL']
			fnt_cells[8].text = font['UniqueID']
			fnt_cells[9].text = font['Version']
			fnt_cells[10].text = font['Manufacturer']
			fnt_cells[11].text = font['Designer']
			fnt_cells[12].text = font['DesignerURL']
			fnt_cells[13].text = font['Vendor']
			fnt_cells[14].text = font['VendorURL']
			fnt_cells[15].text = font['VendorID']
			fnt_cells[16].text = font['Misc1']
		except:
			prepare_doc_table(fnt_cells, font, True)
	else:
		try:
			fnt_cells[0].text = filter(lambda x: x in string.printable, font['FontFileName'])
			fnt_cells[1].text = filter(lambda x: x in string.printable, font['Name'])
			fnt_cells[2].text = filter(lambda x: x in string.printable, font['Family'])
			fnt_cells[3].text = filter(lambda x: x in string.printable, font['Subfamily'])
			fnt_cells[4].text = filter(lambda x: x in string.printable, font['Copyright'])
			fnt_cells[5].text = filter(lambda x: x in string.printable, font['Trademark'])
			fnt_cells[6].text = filter(lambda x: x in string.printable, font['LicenseDesc'])
			fnt_cells[7].text = filter(lambda x: x in string.printable, font['LicenseURL'])
			fnt_cells[8].text = filter(lambda x: x in string.printable, font['UniqueID'])
			fnt_cells[9].text = filter(lambda x: x in string.printable, font['Version'])
			fnt_cells[10].text = filter(lambda x: x in string.printable, font['Manufacturer'])
			fnt_cells[11].text = filter(lambda x: x in string.printable, font['Designer'])
			fnt_cells[12].text = filter(lambda x: x in string.printable, font['DesignerURL'])
			fnt_cells[13].text = filter(lambda x: x in string.printable, font['Vendor'])
			fnt_cells[14].text = filter(lambda x: x in string.printable, font['VendorURL'])
			fnt_cells[15].text = filter(lambda x: x in string.printable, font['VendorID'])
			fnt_cells[16].text = filter(lambda x: x in string.printable, font['Misc1'])
		except:
			pass

def write_android_apps(android_apps, document):
	for app in android_apps:
		document.add_heading(app, level=2)
		p = document.add_paragraph('Developer: ', style='List Bullet')
		p.add_run(android_apps[app][0])
		p = document.add_paragraph('Web URL: ', style='List Bullet')
		p.add_run(android_apps[app][1])
		document.add_paragraph("Embedded Font Software", style='List Bullet')
		for font in android_apps[app][2]:
			if 'AR/VR' in font['UseCase']:
				document.add_paragraph('AR/VR Font - ' + font['Name'], style='List Bullet 2')
			else:
				document.add_paragraph(font['Name'], style='List Bullet 2')
		document.add_paragraph('Screenshot of Font Software Embedded Within Mobile Application', style='List Bullet')
		for path in set(android_apps[app][3]):
			document.add_paragraph(path)
		document.add_paragraph('App Version History', style='List Bullet')
		app_annie_url = 'https://www.appannie.com/apps/google-play/app/' + android_apps[app][2][0][
			'Title'] + '/details/'
		document.add_paragraph(app_annie_url)
		# check levels of enforcement
		document.add_paragraph('Internal Notes', style='List Bullet')
		enf_MT = False
		enf_3p = False
		for font in android_apps[app][2]:
			if font['Enforce'] == 'Yes':
				enf_MT = True
			elif font['Enforce'] != 'Yes':
				enf_3p = True
		if enf_MT:
			document.add_paragraph("Embedded Monotype Font Software", style='List Bullet 2')
			for font in android_apps[app][2]:
				if font['Enforce'] == 'Yes':
					document.add_paragraph(font['Name'], style='List Bullet 3')
		if enf_3p:
			document.add_paragraph("Other Embedded Font Software (UFDA, MyFonts Agreement, Other)",
								   style='List Bullet 2')
			for font in android_apps[app][2]:
				if font['Enforce'] != 'Yes':
					document.add_paragraph(font['Name'], style='List Bullet 3')

def add_picture_in_document_if_exists(document, dirpath, website, font, filename):
	web_research_path_arr = font['Path'].split('/') if(font['Path']) else os.path.join(dirpath, '../web_research', website).split('/')
	# print(web_research_path_arr)
	if(font['Path']):
		web_research_path_arr.pop()
	web_research_path = '/'.join(web_research_path_arr)
	final_image_path = web_research_path + '/images/' + filename
	print(final_image_path)
	if os.path.exists(final_image_path):
		document.add_picture(final_image_path, width=Inches(5.0))

def insert_website_trafic_image(document, font, dirpath, website):
	add_picture_in_document_if_exists(document, dirpath, website, font, 'websiteTrafic.jpg')

def insert_fonts_list_image(document, font, dirpath, website):
	add_picture_in_document_if_exists(document, dirpath, website, font, 'fontsList.jpg')

def write_websites(client, document, oldestURLs, dirpath, pre_csv_path, websites, usecase):
	if usecase == Usecase.DigitalAds:
		document.add_heading(client + ' HTML Ads', level=1)
	elif usecase == Usecase.Website:
		document.add_heading(client + ' Websites', level=1)
	elif usecase == Usecase.PDF:
		document.add_heading(client + ' PDFs', level=1)

	for website in websites:
		document.add_heading(website, level=2)
		if usecase == Usecase.Website:
			document.add_paragraph('Web Fonts', style='List Bullet')
		else:
			document.add_paragraph('Fonts', style='List Bullet')
		for font in websites[website]:
			if 'AR/VR' in font['UseCase']:
				document.add_paragraph('AR/VR Font - ' + font['Name'], style='List Bullet 2')
			else:
				document.add_paragraph(font['Name'], style='List Bullet 2')
		document.add_paragraph('Estimated Website Traffic', style='List Bullet')
		document.add_paragraph('')
		insert_website_trafic_image(document, font, dirpath, website)

		document.add_paragraph('Screenshot of Web Font Software in Use on Website', style='List Bullet')
		insert_fonts_list_image(document, font, dirpath, website)

		displayed_domains = []
		for font in websites[website]:
			if font['VerifiedDomain'] not in displayed_domains:
				displayed_domains.append(font['VerifiedDomain'])
				document.add_paragraph(font['VerifiedDomain'])
		if usecase == Usecase.Website:
			document.add_paragraph('Screenshot of [] Web Font Software in Use on Website to Edit Text', style='List Bullet')
			document.add_paragraph('')
			document.add_paragraph('Screenshot of Web Font Software in Use on Website Dating Back to []',
							   style='List Bullet')
			if pre_csv_path != '' and website in oldestURLs:
				document.add_paragraph(oldestURLs[website]['OldestURL'])

				document.add_paragraph("Current font usage Dating Back to []", style='List Bullet 2')
				for font in oldestURLs[website]['CurrentFonts']:
					document.add_paragraph(font + " - " + oldestURLs[website]['CurrentFonts'][font] , style='List Bullet 3')
		document.add_paragraph('')
		document.add_paragraph('Internal Notes & License Indicators', style='List Bullet')
		# check levels of enforcement
		enf_MT = False
		enf_3p = False
		for font in websites[website]:
			if font['Enforce'] == 'Yes':
				enf_MT = True
			elif font['Enforce'] != 'Yes':
				enf_3p = True
		if enf_MT:
			document.add_paragraph("Monotype Web Font Software", style='List Bullet 2')
			for font in websites[website]:
				if font['Enforce'] == 'Yes':
					document.add_paragraph(font['Name'], style='List Bullet 3')
		if enf_3p:
			document.add_paragraph("Other Web Font Software (UFDA, MyFonts Agreement, Other)", style='List Bullet 2')
			for font in websites[website]:
				if font['Enforce'] != 'Yes':
					document.add_paragraph(font['Name'], style='List Bullet 3')
		li_present = False
		for font in websites[website]:
			if font['LicenseIndicator'] != '' or 'com.myfonts' in font['UniqueID']:
				li_present = True
		if li_present:
			document.add_paragraph('License Indicators', style='List Bullet 2')
			license_indicators = []
			for font in websites[website]:
				if font['LicenseIndicator'] != '' and font['LicenseIndicator'] not in license_indicators:
					document.add_paragraph(font['LicenseIndicator'], style='List Bullet 3')
					if 'projectid' in font['LicenseIndicator']:
						document.add_paragraph(
							'Type: Monotype WFS. Project ID can be searched in MT Admin to find license details')
					elif 'fontids' in font['LicenseIndicator']:
						document.add_paragraph(
							'Type: Linotype.com WF Purchase. Reach out to Monotype CS (cc_internal@monotype.com) for help checking for license details')
					license_indicators.append(font['LicenseIndicator'])
				if 'com.myfonts' in font['UniqueID']:
					li = 'MFWF. ID: ' + font['UniqueID'].rpartition('.')[-1]
					if li not in license_indicators:
						document.add_paragraph(li, style='List Bullet 3')
						document.add_paragraph(
							'A Pre Sales PSD can be submitted for the 4 digit ID in order to obtain the MyFonts license which generated this web font.')
						license_indicators.append(li)


#def create_ap_web(csv_path, dirpath, *nonuse_sites):
def create_ap_web(csv_path, dirpath):
	fonts_list = []
	with open(csv_path,'r', encoding="utf8") as myFile:
		reader = csv.reader(fix_nulls(myFile))
		rows = list(reader)
	
	for row in rows[1:]:
		fontData = makeFontFromCSV(row)
		fonts_list.append(fontData)
	
	outpath = dirpath + ' Websites (EXT3).csv'
	
	output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))
	#write csv headers
	output.writerow(['Brand/Subsidiary','Domain','Font Name', 'Enforceable IP?','Editable Text?' 'License?', 'Page Views', 'Font Comments', 'Website Comments'])
	
	for font in fonts_list:
		if font['UseCase'] == 'Web Font' or font['UseCase'] == 'Digital Ad':
			output.writerow(['',font['Title'], font['Name'], font['Enforce'],'', '', '', font['Remarks'], ''])
			
	#if nonuse_sites:
	#	for site in nonuse_site:
	#		output.writerow(['',site,'system fonts', 'No', 'No', 'NA','','',''])
			

def create_ap_apps(csv_path, dirpath):
	fonts_list = []
	with open(csv_path,'r', encoding="utf8") as myFile:
		reader = csv.reader(fix_nulls(myFile))
		rows = list(reader)
	
	for row in rows[1:]:
		fontData = makeFontFromCSV(row)
		fonts_list.append(fontData)
	
	outpath = dirpath + ' Apps (EXT2).csv'
	
	output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))
	#write csv headers
	output.writerow(['App Name','Platform','App Status','Font Name', 'Enforceable IP?', 'License?', 'App Initial Release Date', 'App Current Version Date', 'Font Comments', 'Website Comments'])
	
	for font in fonts_list:
		if 'Mobile App' in font['UseCase']:
			output.writerow([font['Title'], font['AppPlatform'], 'Active', font['Name'], font['Enforce'], '', '', '', font['Remarks'], ''])
			

def find_iOS_apps(dev_url,my_country,outpath):
	found_apps = []
	apps = []

	cc_list_iOS = [['ae','United Arab Emirates'],['ag','Antigua and Barbuda'],['ai','Anguilla'],['al','Albania'],['am','Armenia'],['ao','Angola'],['ar','Argentina'],['at','Austria'],['au','Australia'],['az','Azerbaijan'],['bb','Barbados'],['be','Belgium'],['bf','Burkina-Faso'],['bg','Bulgaria'],['bh','Bahrain'],['bj','Benin'],['bm','Bermuda'],['bn','Brunei Darussalam'],['bo','Bolivia'],['br','Brazil'],['bs','Bahamas'],['bt','Bhutan'],['bw','Botswana'],['by','Belarus'],['bz','Belize'],['ca','Canada'],['cg','Democratic Republic of the Congo'],['ch','Switzerland'],['cl','Chile'],['cn','China'],['co','Colombia'],['cr','Costa Rica'],['cv','Cape Verde'],['cy','Cyprus'],['cz','Czech Republic'],['de','Germany'],['dk','Denmark'],['dm','Dominica'],['do','Dominican Republic'],['dz','Algeria'],['ec','Ecuador'],['ee','Estonia'],['eg','Egypt'],['es','Spain'],['fi','Finland'],['fj','Fiji'],['fm','Federated States of Micronesia'],['fr','France'],['gb','Great Britain'],['gd','Grenada'],['gh','Ghana'],['gm','Gambia'],['gr','Greece'],['gt','Guatemala'],['gw','Guinea Bissau'],['gy','Guyana'],['hk','Hong Kong'],['hn','Honduras'],['hr','Croatia'],['hu','Hungaria'],['id','Indonesia'],['ie','Ireland'],['il','Israel'],['in','India'],['is','Iceland'],['it','Italy'],['jm','Jamaica'],['jo','Jordan'],['jp','Japan'],['ke','Kenya'],['kg','Krygyzstan'],['kh','Cambodia'],['kn','Saint Kitts and Nevis'],['kr','South Korea'],['kw','Kuwait'],['ky','Cayman Islands'],['kz','Kazakhstan'],['la','Laos'],['lb','Lebanon'],['lc','Saint Lucia'],['lk','Sri Lanka'],['lr','Liberia'],['lt','Lithuania'],['lu','Luxembourg'],['lv','Latvia'],['md','Moldova'],['mg','Madagascar'],['mk','Macedonia'],['ml','Mali'],['mn','Mongolia'],['mo','Macau'],['mr','Mauritania'],['ms','Montserrat'],['mt','Malta'],['mu','Mauritius'],['mw','Malawi'],['mx','Mexico'],['my','Malaysia'],['mz','Mozambique'],['na','Namibia'],['ne','Niger'],['ng','Nigeria'],['ni','Nicaragua'],['nl','Netherlands'],['no','Norway'],['np','Nepal'],['nz','New Zealand'],['om','Oman'],['pa','Panama'],['pe','Peru'],['pg','Papua New Guinea'],['ph','Philippines'],['pk','Pakistan'],['pl','Poland'],['pt','Portugal'],['pw','Palau'],['py','Paraguay'],['qa','Qatar'],['ro','Romania'],['ru','Russia'],['sa','Saudi Arabia'],['sb','Soloman Islands'],['sc','Seychelles'],['se','Sweden'],['sg','Singapore'],['si','Slovenia'],['sk','Slovakia'],['sl','Sierra Leone'],['sn','Senegal'],['sr','Suriname'],['st','Sao Tome e Principe'],['sv','El Salvador'],['sz','Swaziland'],['tc','Turks and Caicos Islands'],['td','Chad'],['th','Thailand'],['tj','Tajikistan'],['tm','Turkmenistan'],['tn','Tunisia'],['tr','Turkey'],['tt','Republic of Trinidad and Tobago'],['tw','Taiwan'],['tz','Tanzania'],['ua','Ukraine'],['ug','Uganda'],['us','United States'],['uy','Uruguay'],['uz','Uzbekistan'],['vc','Saint Vincent and the Grenadines'],['ve','Venezuela'],['vg','British Virgin Islands'],['vn','Vietnam'],['ye','Yemen'],['za','South Africa'],['zw','Zimbabwe']]
    
	for cc in cc_list_iOS:
		url = urlopen(dev_url.replace("/developer/","/" + str(cc[0]) + "/developer/"))
		soup = BeautifulSoup(url.read())
		
		for app in soup.find_all('a', {'class' : re.compile('we-lockup *')}):
			app_url = app['href']
			for foo in app.find_all('div', class_='we-lockup__title'):
				app_name = remove_control_characters(foo.text.strip())
			app_id = app_url.split('/')[-1]
			if app_id not in found_apps:
				apps.append(["iOS",app_name,app_url,app_id,[cc[1]]])
				found_apps.append(app_id)
            #already found this app, add country to the list of stores
			else:
				for ipa in apps:
					if ipa[3] == app_id:
						ipa[4].append(cc[1])
	outpath = outpath + 'iOS_countries.csv'
	output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))
	
	for app in apps:
		found = False
		for region in app[4]:
			if my_country == region:
				found = True
		if found:
			output.writerow([app[0],app[1],app[2],'Available from your local app store'])
		else:
			output.writerow([app[0],app[1],app[2],app[4]])
	
	print('Done! CSV can be found at ' + outpath)

def getIOSApps(fonts_list):
	#iOS_apps: [Title, Developer, Web URL, [fonts], [paths]]
	iOS_apps = []
	app_titles = []
		
	for font in fonts_list:
		if font['AppPlatform'] == 'iOS' and font['Title'] not in app_titles:
			app_titles.append(font['Title'])
	
	app_developer = ''
	app_weburl = ''
	
	if len(app_titles) > 0:
		for app in app_titles:
			#app_fonts_check holds unique fonts by filename		
			app_fonts_check = []
			app_fonts = []
			app_fonts_paths = []
			for font in fonts_list:
				if font['AppPlatform'] == 'iOS' and font['Title'] == app and font['FontFileName'] not in app_fonts_check:
					app_developer = font['AppDev']
					app_weburl = font['AppURL']
					app_fonts.append(font)
					app_fonts_check.append(font['FontFileName'])
					app_fonts_paths.append(str(os.path.dirname(font['Path'])))		
			iOS_apps.append([app,app_developer,app_weburl,app_fonts,app_fonts_paths])

	return iOS_apps

def getAndroidApps(fonts_list):
	#android_apps: [Title, Developer, Web URL, [fonts], [paths]]
	android_apps = {}
	amazon_apps = {}

	for font in fonts_list:
		if font['AppPlatform'] == 'Android':
			if 'Amazon' in font['UseCase']:
				dict = amazon_apps
			else:
				dict = android_apps
			if font['Title'] in dict:
				if font['FontFileName'] not in dict[font['Title']][4]:
					dict[font['Title']][2].append(font)
					dict[font['Title']][3].append(str(os.path.dirname(font['Path'])))
					dict[font['Title']][4].append(font['FontFileName'])
			else:
				dict[font['Title']] = [font['AppDev'], font['AppURL'], [font], [str(os.path.dirname(font['Path']))], [font['FontFileName']]]

	return android_apps, amazon_apps
	
def getWebsites(fonts_list, usecase):
	#websites: [Title, [fonts]]
	websites = {}

	for font in fonts_list:
		if font['UseCase'] == usecase:
			if font['Title'] not in websites:
				websites[font['Title']] = []
			if font not in websites[font['Title']]:
				websites[font['Title']].append(font)
		
	return websites
	
#return list of fonts as well as websites using no fonts per scraper
def readScrapeResults(scrape_results_path):
	scrape_data = []
	fonts_list = []
	websites_no_wf = []
	redirectedURLs = {}
	
	with open(scrape_results_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
	
		for row in reader:
			scrape_data.append(row)
			
	for row in scrape_data[1:]:
		redirectedURLs[row[0]] = row[4]
		#filter out errors from scrape results
		if row[1] =='':
			#filter out No WF sites
			if row[2] !='':
				#font should either be readable or a MT-hosted WF
				fontData = initFontData()
				
				#check for MT-hosted WF (limited info is available for such a font)
				if 'fast.fonts.net' in row[5]:
					fontData['Name'] = row[5][30:72]
					fontData['Family'] = 'Monotype WFS hosted web font'
					fontData['Foundry'] = 'Monotype WFS'
					fontData['Enforce'] = 'Yes - MT WFS - Check Subscription Details'
					fontData['WhyEnforce'] = 'fast.fonts.net in row[5]'
					fontData['VerifiedDomain'] = row[2]
					fontData['Title'] = row[0]
					fontData['UseCase'] = 'Web Font'
					remarks = ['None']
				#not a MT font, full data available
				else:
					filename = row[5].split('/')[-1]
					if 'typekit' in row[5]:
						if 'FontShop' in row[6]:
							fontData['Name'] = 'Typekit hosted FontShop web font: ' + str(row[9])
						elif len(row[10]) > 2:
							fontData['Name'] = row[10]
						else:
							fontData['Name'] = 'Typekit hosted web font'
					elif row[10] == '.':
						fontData['Name'] = filename
					elif 'FontShop' in row[6]:
						if len(row[10]) == 27:
							fontData['Name'] = row[9]
						elif row[9] == 'Web FontFont' and len(row[10]) == 6:
							fontData['Name'] = row[13].split(' ')[0]
						else:
							fontData['Name'] = filename
					elif row[10] != '':
						fontData['Name'] = row[10]
					else:
						fontData['Name'] = filename
					
					fontData['FontFileName'] = filename
					fontData['Copyright'] = row[6]
					fontData['Trademark'] = row[13]
					fontData['Family'] = row[7]
					fontData['Subfamily'] = row[8]
					fontData['LicenseDesc'] = row[19]
					fontData['LicenseURL'] = row[20]
					fontData['UniqueID'] = row[9]
					fontData['Misc1'] = row[29]
					fontData['Version'] = row[11]
					fontData['Manufacturer'] = row[14]
					fontData['Designer'] = row[15]
					fontData['DesignerURL'] = row[18]
					fontData['VendorURL'] = row[17]
					fontData['WebPath'] = row[5]
					fontData['UseCase'] = 'Web Font'
					fontData['Title'] = row[0]
					fontData['VerifiedDomain'] = row[2]
					fontData['LicenseIndicator'] = row[3]

					if 'cloudfront.net' in fontData['Title']:
						fontData['UseCase'] = 'Digital Ad'
				
					checkFontFoundry(fontData)
					#remarks will revert 'Enforce' key if license indicator is present
					fontData['Remarks'] = getRemarks(fontData)
								
				fonts_list.append(fontData)	
			else:
				websites_no_wf.append(row[0])
				
	return [fonts_list,	websites_no_wf, redirectedURLs]
	

def getUseCase(font_path):
	p = Path(font_path)
	use_case = ''
	for i in range(len(p.parts)):
		if 'app_research' in p.parts[i]:
			use_case = 'Mobile App'
			break
		elif 'web_research' in p.parts[i]:
			use_case = 'Web Font'
			break
		elif 'ad_research' in p.parts[i]:
			use_case = 'Digital Ad'
			break
		elif 'pdf_research' in p.parts[i]:
			use_case = 'PDF'
			break
		elif 'amazon_research' in p.parts[i]:
			use_case = 'Mobile App Amazon'
			break
	return use_case

#def validateScrape(input_urls, export_data):
def getScrapeErrors(scrape_results_path):
	
	scrape_data = []
	scrape_errors = []
	
	with open(scrape_results_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		
		for row in reader:
			scrape_data.append(row)
	
	#remove header and subset for uniqueness
	scrape_data = scrape_data[1:]
	scrape_data = [list(x) for x in set(tuple(x) for x in scrape_data)]
	
	for row in scrape_data[1:]:
		if row[1] != '':
			scrape_errors.append([row[0],row[1]])
	
	scrape_errors = [list(x) for x in set(tuple(x) for x in scrape_errors)]
	
	return scrape_errors

def checkMissedWebsites(input_websites_path, scrape_results_path):
	input_websites = []
	scrape_websites = []
	missed_websites = []
	input_websites_clean = []
	
	with open(input_websites_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		for row in reader:
			input_websites.append(str(row)[2:-2])
	
	with open(scrape_results_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		for row in reader:
			scrape_websites.append(str(row[0]))	

	#regex to replace 'http://www.' etc to match form of web scrape output
	#may need to check separately for regex5 pattern as this could create false positives for websites such as e.g. 'http://www.mystorewww.com'
	regex1 = re.compile('https://www.')
	regex2 = re.compile('http://www.')
	regex3 = re.compile('https://')
	regex4 = re.compile('http://')
	regex5 = re.compile('www.')
	
	for input_website in input_websites:
		input_website = re.sub(regex1, '', input_website)
		input_website = re.sub(regex2, '', input_website)
		input_website = re.sub(regex3, '', input_website)
		input_website = re.sub(regex4, '', input_website)
		input_website = input_website.split('/',1)[0]
		input_websites_clean.append(input_website)	
		
		
	for input_website in input_websites_clean:
		missed = True
		for scrape_website in scrape_websites:
			if input_website in scrape_website:
				missed = False
		if missed:
			missed_websites.append(input_website)
	
	#remove duplicates
	missed_websites = set(missed_websites)
	
	return missed_websites

def getRedirects(scrape_results_path):
	#2 types of redirects- website redirects and the scraper crawling outside the provided domain
	scraped_websites = []
	redirects = []
	redirects.append(['',''])
	redirects_conf = []

	with open(scrape_results_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		for row in reader:
			scraped_websites.append(row)	

	for website in scraped_websites[1:]:
		website_redir = [str(website[0]), str(website[4]), 'redir']
		website_crawl = [str(website[0]), str(website[2]), 'crawl']
		if len(website_redir[1]) >0 and website_redir[0] not in website_redir[1] and website_redir not in redirects:
			redirects.append(website_redir)
		elif len(website_crawl[1]) > 0 and website_crawl[0] not in website_crawl[1] and website_crawl not in redirects:
			redirects.append(website_crawl)

	redirects = redirects[1:]
		
	#create checklist of redirects for confirmation
	t = []
	v = []
	for redirect in redirects:
		t.append(redirect[0] + ' || ' + redirect[1])
		v.append(True)

	my_dialog = dialog.Dialog(title = 'Verify Redirects ' + SCRIPT_VERSION, cancellable = True)
	cl = dialog.Checklist(titles = t, values = v, height = 500)
	my_dialog.add(cl)

	if dialog.show(my_dialog):
		redirects_conf = list(zip(redirects, cl.values))
	
		outpath = scrape_results_path.split('.',-1)[0] + '_validated.csv'
		output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))

		for website in scraped_websites:
			flag = True
			website_check1 = [str(website[0]), str(website[4])]
			website_check2 = [str(website[0]), str(website[2])]
			for redir in redirects_conf:
				redir_check = [redir[0][0], redir[0][1]]
				redir_type = redir[0][2]
				if website_check1 == redir_check and redir[1] == False:
					flag = False
				elif website_check2 == redir_check and redir[1] == False:
					flag = False
				elif website_check1 == redir_check and redir[1] == True and redir_type == 'redir':
					website[0] = website[4]
					flag = False
					output.writerow(website)
				elif website_check2 == redir_check and redir[1] == True and redir_type == 'crawl':
					website[0] = website[2]
					flag = False
					output.writerow(website)
			if flag:
				output.writerow(website)
	else:
		outpath = scrape_results_path.split('.',-1)[0] + '_redirects.csv'
		output = csv.writer(open(outpath,'w',encoding='utf8', newline = ''))
		output.writerow(['Input Domain','Scraped Domain'])
		for redirect in redirects:
			output.writerow([redirects[0],redirects[1]])
		print('Check list canceled. A CSV file of redirects has been saved here: ' + outpath)

'''
def RemoveRedirects(scrape_results_path, validated_redirects_path):
	scraped_websites = []
	redirects = []
	with open(scrape_results_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		for row in reader:
			scraped_websites.append(row)
	with open(validated_redirects_path, newline='',encoding='utf-8') as f:
		reader = csv.reader(f)
		for row in reader:
			redirects.append(row)
	for website in scraped_websites:
			flag = True
			website_check1 = [str(website[0]), str(website[4])]
			website_check2 = [str(website[0]), str(website[2])]
			for redirect in redirects:
				redir_check = [redir[0][0], redir[0][1]]
				redir_type = redir[0][2]
				if website_check1 == redir_check and redir[1] == False:
					flag = False
				elif website_check2 == redir_check and redir[1] == False:
					flag = False
				elif website_check1 == redir_check and redir[1] == True and redir_type == 'redir':
					website[0] = website[4]
					flag = False
					output.writerow(website)
				elif website_check2 == redir_check and redir[1] == True and redir_type == 'crawl':
					website[0] = website[2]
					flag = False
					output.writerow(website)
			if flag:
				output.writerow(website)
'''		
def getRemarks(fontData):
	remarks = ['None']
	
	#Hosted Services
	if 'WebPath' in fontData:
		if 'parastorage' in fontData['WebPath']:
			remarks.append('Wix')
			fontData['Enforce'] = 'No'
		elif 'shopifycdn' in fontData['WebPath']:
			remarks.append('Shopify')
			fontData['Enforce'] = 'No'
		elif 'successfactors.com' in fontData['WebPath']:
			remarks.append('Success Factors Portal Customer')
			fontData['Enforce'] = 'Check with Legal'
		elif 'selectminds.com' in fontData['WebPath']:
			remarks.append('Oracle - Select Minds Portal Customer')
			fontData['Enforce'] = 'No'
		elif 'typekit.net' in fontData['WebPath']:
			remarks.append('Typekit')
			fontData['Enforce'] = 'No'
		elif 'gstatic' in fontData['WebPath']:
			remarks.append('Google-Hosted')
			fontData['Enforce'] = 'No'
		elif 'fast.fonts.net' in fontData['WebPath']:
			remarks.append('MT WFS')
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'fast.fonts.net in WebPath'
		
	#MyFonts Web Font, provides ID for license lookup
	if 'com.myfonts' in fontData['UniqueID']:
		MFWF = 'MFWF. ID: ' + fontData['UniqueID'].rpartition('.')[-1]
		remarks.append(MFWF)
	
	#font muckery/tomfoolery
	if 'fontsquirrel' in fontData['Misc1'].lower() or 'fontsquirrel' in fontData['Misc2'].lower():
		remarks.append('FontSquirrel')
		#font muckery/tomfoolery
	if 'font squirrel' in fontData['Misc1'].lower() or 'fontsquirrel' in fontData['Misc2'].lower():
		remarks.append('FontSquirrel')
		
	if fontData['VendorID'] != '':
		if fontData['VendorID'].decode('utf-8') == 'MACR':
			remarks.append('Macromedia Fontographer')
		elif fontData['VendorID'].decode('utf-8') == 'ATEC':
			remarks.append('ALLTYPE')

	if 'ufonts_' in fontData['FontFileName']:
		remarks.append('ufonts')
	
	#Open Source fonts & Other confusing one-offs
	if ('google' in fontData['Trademark'].lower() or 'google' in  fontData['Copyright'].lower()):
		remarks.append('Google Font')
	elif 'typekit' in fontData['LicenseURL'].lower():
		remarks.append('TypeKit in LicenseURL')
	elif 'SIL Open Font License' in fontData['LicenseDesc']:
		remarks.append('SIL OFL')
	elif 'Apache License' in fontData['LicenseDesc']:
		remarks.append('Apache OFL')
	
	if ('bitstream vera' in fontData['Copyright'].lower()) or ('bitstream vera' in fontData['Trademark'].lower()) or ('bitstream vera' in fontData['Name'].lower()): 
		remarks.append('bitstream vera')
		
	elif ('arimo' in fontData['Copyright'].lower()) or ('arimo' in fontData['Trademark'].lower()) or ('arimo' in fontData['Name'].lower()): 
		remarks.append('arimo')
	
	elif ('dejavu' in fontData['Copyright'].lower()) or ('dejavu' in fontData['Trademark'].lower()) or ('dejavu' in fontData['Name'].lower()): 
		remarks.append('dejavu')
	
	elif ('deja vu' in fontData['Copyright'].lower()) or ('deja vu' in fontData['Trademark'].lower()) or ('deja vu' in fontData['Name'].lower()): 
		remarks.append('dejavu')
	
	elif ('fontstruct' in fontData['LicenseDesc'].lower()) or ('fontstruct' in fontData['Trademark'].lower()) or ('fontstruct' in fontData['Manufacturer'].lower()): 
		remarks.append('fontstruct')

	
	#System fonts
	#MacOS
	pattern = '[0-9]{1,2}\.[0-9][a-z][0-9]{1,2}[a-z][0-9]{1,2}'
	
	if len(re.findall(pattern,fontData['UniqueID']))>0:
		remarks.append('MacOS')
	
	if len(remarks)>1:
		remarks = remarks[1:]
	
	return remarks
	


def checkFontFoundry(fontData):
	
	checkMonotype(fontData)
	if fontData['Foundry'] == '' and fontData['Enforce'] == '':
		checkUFDA(fontData)
	#make sure to check here for top marks
	if fontData['Foundry'] == '' and fontData['Enforce'] == '':
		checkTopMarks(fontData)
	if fontData['Foundry'] == '' and fontData['Enforce'] == '':
		fontData['Enforce'] = 'No'
	
	checkMyFonts(fontData)
		
#input: fontData Dictionary
#output: boolean for whether the font is likely enforceable IP
def checkMonotype(fontData):
	#isMTIP = False
	#mt_foundry = ''
	#Looks for presence of a manufacture year in the copyright info
	house_marks = ['mt','ff','lt','bt','itc','icg']
	
	yearBool = False
	try:
		if len(fontData['Copyright']) > 0:
			year = [int(s) for s in fontData['Copyright'].split() if s.isdigit()]
			year = [x for x in year if x in range(1980,2020)]
			if len(year)>0:
				yearBool = True
	except(KeyError):
		pass

	#Check for MT foundry reference in metadata
	#Monotype
	#The Monotype Corporation, 'MT'
	#FontShop aka FontFont, 'FF'
	#Linotype, 'LT'
	#Bitstream, 'BT'
	#International typeface corporation, 'ITC'
	#Image Club, 'ICG'
	#Ascender
	#MT_foundries = ['bitstream', 'monotype', 'linotype', 'fontshop', 'fsi fonts', 'font shop', 'image club', 'international typeface corporation','fontfont','Ascender']
	
	
	#for foundry in MT_foundries:
	#	if (foundry in fontData['Copyright'].lower()) or (foundry in fontData['Trademark'].lower()) or (foundry in fontData['UniqueID'].lower()):
	#		mt_foundry = foundry
	
	fieldsToCheck = ['Copyright','Trademark', 'UniqueID', 'Designer']
	for field in fieldsToCheck:
		font_identifier = fontData[field].lower()
		if 'bitstream' in font_identifier:
			fontData['Foundry'] = 'Bitstream'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'Bitstream in ' + field
			break
		elif 'monotype' in font_identifier:
			fontData['Foundry'] = 'Monotype'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'Monotype in ' + field
			break
		elif 'linotype' in font_identifier:
			fontData['Foundry'] = 'Linotype'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'Linotype in ' + field
			break
		elif 'fontshop' in font_identifier:
			fontData['Foundry'] = 'FontShop'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'FontShop in ' + field
			break
		elif 'fsi fonts' in font_identifier:
			fontData['Foundry'] = 'FontShop'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'fsi fonts in ' + field
			break
		elif fontData['VendorID'] == 'FSI':
			fontData['Foundry'] = 'FontShop'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'FSI in VendorID'
			break
		elif 'image club' in font_identifier:
			fontData['Foundry'] = 'Image Club Graphics (Monotype Library)'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'image club in ' + field
			break
		elif 'international typeface corporation' in font_identifier:
			fontData['Foundry'] = 'Monotype ITC'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'international typeface corporation in ' + field
			break
		elif 'fontfont' in font_identifier:
			fontData['Foundry'] = 'FontShop'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'fontfont in ' + field
			break
		elif 'ascender' in font_identifier:
			fontData['Foundry'] = 'Ascender'
			fontData['Enforce'] = 'Yes'
			fontData['WhyEnforce'] = 'ascender in ' + field
			break
	
	#Filter: check for nearly-always-enforceable-by-mark fonts 
	#topMarks = ['helvetica','gill sans', 'gillsans', 'soho gothic', 'neo sans', 'trade gothic']
	
	#if not isMTIP:
	#	for topFont in topFonts:
	#		if topFont in fontData['Name'].lower():
	#			isMTIP = True
	#			mt_foundry = lookUpTM(topMark[0])
	
	#Filter: Adobe
	if fontData['Foundry'] == '' and fontData['Enforce'] == '':
		if 'adobe systems' in fontData['Copyright'].lower():
		#check for new fonts first,
			if yearBool:
				nonMT_years = ['2004','2005','2006','2007','2008','2009','2010','2011','2012','2013','2014','2015','2016','2017','2018','2019','2020','2021','2022']
				for yr in year:
					if yr not in nonMT_years:
						#check for house marks
						for mark in house_marks:
							if mark in fontData['Name'].lower():
								fontData['Foundry'] = 'Adobe'
								fontData['Enforce'] = 'Check with Legal'
	
	#Filter:Esselte-Letraset
	if fontData['Foundry'] == '' and yearBool and fontData['Enforce'] == '':
		es_let = ['esselte','letraset']
		for foundry in es_let:
			if (foundry in fontData['Copyright'].lower()) or (foundry in fontData['Trademark'].lower()):
				for yr in year:
					if yr < 2000:
						fontData['Foundry'] = 'Monotype ITC (Esselte-Letraset)'
						fontData['Enforce'] = 'Yes'
						fontData['WhyEnforce'] = 'Monotype ITC (Esselte-Letraset)'
						

	#Filter:nonMT Fonts
	#	nonMT Fonts:
	#		reference to Font Struct -> older font creation software, not MT font
	#		non-commercial Ascender, BT, MT, etc fonts
	nonMT_fonts = ['google','texgyretermes','paratype', 'Dubai','bitstream vera','bebas kai','bebas neue','bebasneue','open sans', 'arimo','tinos','dyslexic','dejavu', 'deja vu', 'droid sans', 'droidsans', 'droidmono', 'droid mono', 'droidserif', 'droid serif','droidarabic','droid arabic', 'droid kufi','droidkufi', 'liberation sans', 'liberation serif', 'liberationserif', 'liberationsans','fontstruct', 'font struct', 'noto sans', 'noto serif', 'noto', 'sap72', 'sap 72', 'jiotype','rn house sans','rnhousesans','segoe','georgia','comic sans','verdana','tahoma','calibri','raleway','Raleway','Libre Baskerville','LibreBaskerville','librebaskerville','Sanchez','CevicheOne','Ceviche One','Cecivhe']
	
	for font in nonMT_fonts:
		if (font.lower() in fontData['Copyright'].lower()) or (font.lower() in fontData['Trademark'].lower()) or (font.lower() in fontData['Name'].lower()):
			fontData['Foundry'] = ''
			fontData['Enforce'] = 'No'
			fontData['WhyEnforce'] = ''
			return

	if 'typekit' in fontData['LicenseURL'].lower():
		fontData['Enforce'] = 'No'
		fontData['WhyEnforce'] = ''
		return
	
	#filter may catch too much, check again for known exceptions to the exceptions
	
	nonMT_fonts_exceptions = ['segoe print', 'segoeprint']
	fieldsToCheck = ['Copyright', 'Trademark', 'Name']
	for font in nonMT_fonts_exceptions:
		for field in fieldsToCheck:
			if (font in fontData[field].lower()):
				fontData['Foundry'] = 'The Monotype Corporation'
				fontData['Enforce'] = 'Yes'
				fontData['WhyEnforce'] = font + ' in ' + field

#looks for com.myfonts.... in the Unique ID string
def checkMyFonts(fontData):
	if 'com.myfonts' in fontData['UniqueID'].lower():
		fontData['Enforce'] = 'Yes- MyFonts Customer - Review Existing License'
		fontData['WhyEnforce'] = 'Yes- MyFonts Customer - Review Existing License'

#need to add 'setup','writings' foundries, and review 'cast'
def checkUFDA(fontData):
	fieldsToCheck = ['Copyright', 'Trademark', 'UniqueID', 'Designer']
	ufda_foundries = ['1871 Project', 'Akarija Studio', 'AKTF','Alex Meier', 'Almazova Dolzhenko', 'Anara Ara', 'Anti.biz', 'AravisFonts','ArlindType','Artecture','AsbakLab','Andanistas','Arttype7', 'Astageni','Atasi Studio','Azzam Ridhamalik','Blancoletters','Bogstav','Brainware Graphic', 'Carbon Milk Fonts',' CAST ','Character Type', 'Cindy Kinash', 'Connary Fagen','Corentin Noyer','Corradine Fonts', \
						'Creative Grenade','Counterpoint Type Studio', 'Cultivated Mind','Daldsgh','Daylight Fonts','Dee Creations','Deeezy','Demockups','Denis Ignatov','DePlictis Type','DesignClusters','dgs designs','Diego Massaro','Digitype Studio','Dixie Type Co','Dmitry Bogolyubov','Dora Typefoundry','Dropper','Drysk Creative','DS Type', 'DSTYPE','Dustbunnies Everywhere','EdricStudio', 'Elswick','Elyas Beria','Emile Lefebvre','Encolab', 'Ensotype','Ephemera Fonts','Epiclinez','Estudio Ponce Conteras','Factory 738','Fando Fonts','Fenotype','Fiqiart','Five Goods','Fontastica','Fontfabric','FSY Creative Lab','Fype Co','FontSmith', 'Gleb Guralnyk', \
						'Good Form','Goodigital13', 'Good Form','Great Indoors', 'Greentypestudio','Grigorij Gushchin','GRIN3 (Nowak)','Gror','Guintype','Gurup Stüdyo','Halymunt Studio', 'Hanken Design Co.', 'Hanoded','Harnessless Type','Helenas Letter','Henta Zkhr','Hikhcreative ', 'Hoopoe','Hurme','Hustle Supply','"HVF Studio','Idesain Creative','Increments','Intelligent Design', \
						'Irina Mir', 'IUnderline Art', 'Jasonkovac','Jeff Levine', 'Jeremie Gauthier','Joebob Graphics','Jonahfonts', 'Jordyn Alison Designs','JprintStudio','Jujumisur’s Ficus','Jukebox Collection', 'Kaer','Kazer Studio','Kellie Jayne Studio','Kuldd','Miguel & Daniel HernÃƒÂ¡ndez','Miguel&DanielHernndez','Daniel Hernández','Latinotype','Laura Worthington','Lella7','LevBez','Lewis McGuffie Type','Liartgraphic','Lloyd David Designs','Lone Army','Los Andes','Mabhal Studio','Magnum Studio','Majed ps','Marsnev','Maryam Kamal Co','Maxim Suvoroff','Memoodesigns','Mevstory Studio','Michael Prewitt','Missin Glyphs','MIX.Jpg','Molly Suber Thorpe','Monoco Type','Moritz Kleinsorge','Muhittin Güneş','Mr. Black Fonts', \
						'Muhittin Güneş', 'NapoleanServices','Niznaztype','NJ Studio','NorFonts','Nun Creative','Nurmiftah','Oakstone Creative','Océane Moutot','OCSstudio','Paavola Type Studio','Papanapa','Papermode Co','Par Défaut','PizzaDude.dk','pizzadude','Plasebo Studio','PotatoGraphick','Pratama Yudha','Rillatype',' RMType ', 'RodrigoTypo', 'Rotsoul','Rui Nogueira','Rumors Foundry',' Samuel Morgan', 'Samuel Vicente Types', 'Scratch Design','Set Sail Studios', 'Showup! Typefoundry','Shriftovik',' Sign Studio ','SilverMintDesign', 'Solidtype','Solotype',
						'Starsoft', 'Stolat Studio', 'SullivanStudio', 'Supernulla', 'Syifidz', 'Taner Ardali', 'Teakword', 'Thanoestd', 'The Paper Town', 'Tipo Pipel', 'Tipo Pèpel', 'Tipotype', 'Titis', 'Tomass Gavars', 'Tony Fahy Font Foundry', 'TSA Creatuve', 'Type Salon', 'Type Trends', 'Type Type', 'Type.write.type', 'TypeType', 'Typoforge','Typographias', 'UP UP Creative', 'VType','VzType', 'Wayne Fearnley', 'Winston Type Co','Wirtu', 'words+pictures', 'Wraith Types', 'Wyarecreatype','Yatanski','Yoon Design','Zamjump','Ziza Type','Zumkin','House Industries','House Industries/Brand Design Co','URW','Fontfabric LLC/Inc.','Exljbris','Jos Buivenga','exljbris','http://fenotype.com','Fenotypefaces','YoonDesign','(주)윤디자인연구소','Zetafonts','Zetafonts','Kostas Bartsokas','DSType Foundry','Dino dos Santos','Pedro Leal','DSType','Mika Melvas','Mika Melvas','mahti.fi','Set Sail Studios Ltd','Hanoded','David Kerkhoff','Kimmy Design','Toni Hurme','www.hurmedesign.fi','Comicraft','My Creative Land','Elena Genova','Aerotype','www.aerotype.com','Miggas','Typoforge Studio','Device','Rian Hughes','devicefonts','Cultivated Mind','TipoType','Paulo Goode','Paulo Goode','JOEBOB graphics','www.joebob.nl','Ahmet Altun','Akarija Studio','akarija.studio@gmail.com','Hegemony','Alina9900','Almazova Dolzhenko','Anara Ara','Anti.biz','AravisFonts','Arukidzfl','Asphodel Design','Astageni','Nur Solikh','Batya','BegArtOne Type','Begiko','Carbon Milk Fonts','Cindy Kinash','Collectype','Connary Fagen','Creative Grenade','Dixie Type Co.','Dusty Lilac Lettering','Elswick','Ensotype','Eternaleyes','FontsRoyale','Frantic Disorder','Good Form','Greentypestudio','Hanken Design Co.','Heako','Hoopoe','House of X','HVF Studio','Inermedia Studio','Irina Mir','IUnderline Art','Jasonkovac','Jehoo Creative','Jeremie Gauthier','Jordyn Alison Designs','Kaer','Kisla','Kiyoung Lee','Latinotype Mexico','Leonor Capricho','LetterCraft Studio','Lithographe','M V Dilanson','MonogramBros','Muykyta','NapoleanServices','OWType','Petterco','PixelHive Pro','PranjTech','RMU','Roman Melikhov','Rotsoul','Samuel Vicente Types','Scratch Design','Setup','SevenType','Showup! Typefoundry','Solidtype','Starsoft','Stolat Studio','SullivanStudio','Supernulla','SweetCake','Swell Type','Syifidz','Taner Ardali','Teakword','Thanoestd','The Paper Town','The Questa Project','Tijs Krammer','Tipo Pèpel','Titis','Tomass Gavars','Tony Fahy Font Foundry','TSA Creatuve','Type Salon','Type Trends','Type.write.type','TypeType','UP UP Creative','Vintage Voyage Design Supply','VType','WAP Type','Wayne Fearnley','Winston Type Co.',' WR ','Wraith Types','Writings','Wyarecreatype','Susan Brand Design','Remedy667','Mix Fonts','Micreative Studio','PrestigeType','Australian Type Foundry','Phat Phonts','Rodrigo Typo','Typerookie','Typekiln','Dino Feed','Muksal Creatives','Umka Type','Inari Type','VP Type','Goicha','ZenTpe','Zenstyle','Lavender','Godzillab','Formatype Foundry','Coconuts Design','Julien Fincker','Gittype','Zeune Type Foundry','Natalja Ameline','Ingrimayne Type','Degarism','dmtype','EllenLuff','Graphic Studio 33','mark simonson','marksimonson','Svetoslav Simov','Svet Simov','Paratype','Rene Bieder','René Bieder','Mostardesign','Type Dynamic','Dharma Type','The Northern Block Ltd','G-Type','Parachute','Typodermic','Adam Ladd','Kastelov','T-26','Jason Vandenberg','SilkType','K-Type','insigne','Stawix','Mans Greback','Julien Saurin (laGoupil.S&C Type)','Wiescher Design','Rsz Type (same as Resistenza)','P22','ogj Type Design','Talbot Type','FSD','Resistenza','Inigo Jerez Quintana - Extratye','Counterpoint Studio','Tegami Type','Trial by Cupcakes','RodrigoTypo','Brink','Colmena','Hafeez UI Haq','OStype','Mint Type','Anita Jürgeleit','Zefrar','Fortunes Co.','Rafaeiro Typeiro','Nihar Mazumdar','Lets Go','Egor Stremousov','Graviton','Kulturrrno','Santi Rey','Taweka','LGF','Peach Creme','Roland Huse Design','Anomali Creative','Frog1812','Aiyari','LafonType','LuxFont','Satori TF','PepiQuokka','Jorsecreative','ATK Studio','Rodrigo Fuenzalida','Alternate Glyphs-cant login','Zet Design','Dima Pole','OrakArik','Letterara','VIP Graphics','Etewut','Tural Alisoy','Cyanotype','Spacetype','HilwArt','FontPeople','Alfareaniy','M.A. Ho-Kane','Fauzistudio','MXMV','Rivian','Albascreative','Parker Creative','Dicky Syafaat','Philipp Richter','Redcollegiya','Abbastype','Hazztype','Type Toucan','BUFC Creative','Giaimefontz','Kotak Kuning Studio','Shapovalov Fonts','Urmytype','Gleb Guralnyk','dmrailabstd','Surplus Type Co','Nootype','Neelatype','AZCRTV Studio','Ayca Atalay','CostaType','Sihan Wu','Zeniac','Studio Yang','Punchform','Moontesk','Grey Fortress','Cretype','Cmeree','Circastudio','Christian Gamga Pedro','Antitype','611 Studio','Keanu','Zealab Fonts Division','Riljs','Will Albin-Clark','Vertigo','Tigade Std','Sipanji21','ScovType','Nordebrink Studio','Ira Natasha','Grezline Studio','Andrey Font Design','Buhasa Type','OneSixOne','2D Typo','Scoothtype','Cao Fila','Mi Chen','Belovestudio','FarahatDesign','Adorae Types','Stefani Letter','Salt Color','Beware of the moose','Alex Jones','RNS Fonts','Macizo.com.mx','NeutroneLabs','DimitriAna','Atelje Altmann','Abbyland','Floodfonts','Fontease','MTC Graphic','Adelina Apostolova','AnkaDrozd','Blacksheep Studio','Folding Type','JK Creative','Keren Studio','Stefan Stoychev','Tropical Type Foundry','Zaza type','Mandarin','Schizotype','Eugene Bunin','Curated texts','Haffastudio','Tom Hallgren','Deals on Design','Prop-a-ganda',' WEP ','Ultravioleta','Pedro Mello Type Foundry','Nurrontype','Nurrehmet','Nantia.co','Milan Pleva','Just Bia','John Misael','Hatftype','Black Studio','Sawdust','ENCI','Stuart Hazley','The Ocean Studio','Yumna Type','Furkan Ilbay','Arabetics','Great Studio','Sweetest Goods','Wahya and Sari Co','NaumType','Din Studio','PabType','Peterdraw','Alphabet Agency','Dikas Studio','Edd\'s Aurebesh Fontworks','AndriyFM','Brithos Type','LuvOMIE','Designpiraten','Oz Design','Letter Omega Typefoundry','Teweka','Menganz','Road to Venice Type','E-phemera','Artyway','CoastalType','SteerFonts','Akifatype','Khaiuns','Latour','Aronetiv','TypoBureau Studio','Nikishlab','ErlosDesign','Octodeca',' SouthEast ','Max Prive','Kufic Studio','Franklin Veiga','Ksenia Belobrova','Balpirick','Aim_CS','Dhan Studio','Mr Typeman','Sara Kok','Sebeningjingga','Neue Radial','Ryoichi Tsunekawa',' Ingo ','Red Rocket Signs',' Nois ','Suza Studio','Lebbad Design','Vitaliy Tsygankov','Typestetic Studio','Suamzu Art','Studio Bayley','SimpleType Studios','Shanu Hu','Samrenal Studio','Rocky Nacpil','Plamen Atanasov','Patria Ari','Minor Praxis','Mega Type','JC Creation Design','Art with Words','In-House International','Esintype','Archetypo.xyz','The Design Speak','O Type Foundry','Studio&Story','BoxTube Labs','Vintage Type Company','Contrafonts','Christian Gruber','Illusletra','words+pictures','Andinistas','Studio K','Pizzadude','Artegra','Senzana','Katrinelly','Jrmuitos','LetterStock','Nadezda Gudeleva','Szymon Furjan','Arkilion','Growcase','Mikolaj Grabowski','EchadType','Andrejs Kirma','Irina Kryvets','SimpleBits','Salih Kizilkaya','Yasemin Varlik','Si47ash Fonts','Drawwwn','MatchaJoeJoe Fonts','Typophobia','Hemphill Studio','Octotypo','Fontsoon','Di-fonts','Oliveira 37','Marianna Orsho','Ludere Studios','Kiwiplates','Hasta Type','HansCo','Fonderia Serena','Fidan Fonts','Blonde Typefoundry','Andresach','Wildstripe','Viaction Type.Co','Unsifonts','Tipogra Fio','Something and Nothing','Limelight Artistry','Gilar Studio','Ettie Kim Studio','Eko Bimantara','Ekahermawan','Borutta Group','MindMooring','Shinntype','Aboutype','Adevio Studio','Autographis','TOMO','Anatoletype','Typedepot','Kimberly Geswein','LetterPerfect','TypeSETit','Club Type','Green Type','Sassoon-Williams','Calligraphics','Alan Meeks','S&C Type','Ana\'s Fonts','Yung & Frish','Typoets','The Tree is Green','StudioJASO','Roman Ilynykh','Leeza Chepugova','Heypentype','FlashGraphics','Dieza Design','Daria Elyasova','CBRTEXT','Carmel Type Co.','Bejeletter','Anonymous Typedesigners','Jorg Schmitt','Graphite','Matt Frost','Pixilate','Typotheticals','Intellecta Design','Letritas','Wilton Foundry','Synthview','North Type','Bomparte\'s Fonts','Hashtag Type','Armasen','Stephen Rapp','Blambot','Eurotypo','Borges Lettering','Juraj Chrastina','WTFont','Gie Studio','Braw Type','Providence Type','StereoType Fonts','JCFonts','Breauhare','Tour De Force','Vanarchiv','Barnbrook Fonts','Pedro Teixeira','Boover Software','Piñata','Blue Vinyl','Lanston Type Co','IHOF','Hamilton Wood Type Collection','Slava Antipov','Nina Belikova','LayarBahtera','Hindia Studio','HeadFirst','Nelson Borhek Press','The Ampersand Forest','Jen Wagner','Heinzel Std','Good Java Studio','Bulent Yuksel','pictomato','Type Fleet','Letterhead Studio VV','Jadugar Design','Tipos Pereira','Jeff Kahn','Type Atelier','Larin Type co','Janworx','R9 Type Design','Beary','Linh Nguyen','Rafael Jordan','Typedifferent','NicolassFonts','NREY','Great Scott','Studio Buchanan','Aga Silva','Handpainted Type','Ruben R dias','Ali Hamidi','sugargliderz','Fontop','Fontdation','Knorke','Koray Ozbey','FlehaType','Burford','Varsity Type','Dirtyline Studio','FMD','Orpin Type','Saffatin','Sulthan Stydui','New Fonts','Tyler Finck','Tarallo Design','Barmoor Foundry','Hackberry Font Foundry','Wild Edge','The Refinery','Joelmaker','Marc Lohner','CreativemediaLab','Kodhibanks','Missy Meyer','Siwox Studios','Bhadalstudio','TypeUnion','Attract Studio','Daggertypo','Underscore','A New Machine','Iam314','Typolis','Nilson Art Design','Fargun Studio','Essque Productions','Majestype','Eliezer Drawe','Pelavin','Zeenesia Studio','Heummdesign','Pepper Type','Talavera','SSI.Scraps','Rekord','Rethink Design Co','CKhans Fonts','Uncurve','FHFont','MMC Typodrome','Martype','Multype Studio','LetterPalette','Ana Prodanovic','Peter Huschka','Fo Da','Eraky','Nasir Udin','Page Studio Graphics','inkstypia','Sketch and Serve','Khoir','Etcsupply','One Line Design','ArtyType','Graphicfresh']

	for field in fieldsToCheck:
		font_identifier = fontData[field].lower()
		if 'house industries' in font_identifier:
			fontData['Foundry'] = 'House Industries'
			fontData['Enforce'] = 'House Industries - Check with Legal'
			fontData['WhyEnforce'] = 'House Industries in ' + field
		elif 'studio k' in font_identifier and 'studio kmzero' not in font_identifier:
			fontData['Foundry'] = 'Studio K'
			fontData['Enforce'] = 'Yes(UFDA)'
			fontData['WhyEnforce'] = 'studio k (not studio kmzero) in ' + field
		else:
			for ufda in ufda_foundries:
				if ufda.lower() in font_identifier:
					fontData['Foundry'] = ufda
					fontData['Enforce'] = 'Yes (UFDA)'
					fontData['WhyEnforce'] = ufda + ' in ' + field

def checkTopMarks(fontData):
	topMarks = ['helvetica','gill sans', 'gillsans', 'soho gothic', 'neo sans', 'trade gothic', 'avant garde gothic', 'frutiger', 'avenir', 'century gothic', 'ff din', 'futura lt', 'itc benguiat', 'itc edwardian script', 'swiss 721', 'linotype didot', 'optima', 'rockwell', 'vag rounded', 'ff mark', 'snell roundhand', 'officina sans','itc eras', 'bradley hand', 'rage italic', 'chiller', 'ff milo']
	
	for topMark in topMarks:
		if topMark in fontData['Name'].lower():
			fontData['Enforce'] = 'Check with Legal'
			fontData['WhyEnforce'] = topMark + ' in Name'
	
	false_positives = ['mcgill sans', 'helveticamazing']
	
	for fp in false_positives:
		if fp in fontData['Name'].lower():
			fontData['Enforce'] = ''
	
		
	
	
#def getFontData
#input: Editor,path
#purpose: extract useful info from Editor and path and return as a list
#output:
#
#
#  FontData & FontUseData(###) 
#	Name: Font Name- Either the Full Name from the NAME table or the file name if Full Name is null 
#	MTIP: Does this appear to be Monotype-enforceable IP?	
#	Remarks
#		System-Type Fonts
#			MacOS font w/ specific version
#			Windows font
#		
#	Copyright
#	Trademark
#	Family
#	Subfamily
#	LicenseDesc
#	LicenseURL
#	UniqueID
#	Misc1: e.g. 'FontSquirrel'
#	Misc2
#	Version
#	Manufacturer
#	Designer
#	DesignerURL
#	Vendor
#	VendorURL
#	VendorID
#	FontFileName
#	FileType
#	Path
#	WebPath
#	Foundry
#	Enforce
###	UseCase: e.g. 'website','app'
###	Title: e.g. Website Domain, App Title
###	AppPlatform
###	AppDeveloper
###	AppURL
#### VerifiedDomain: webscrape- example web page where font use can be documented
#### LicenseIndicator: webscrape- license indicator

def getFontData(font, path, arvrFonts,fileName):
	#returns a dictionary with all values as ''
	fontData = initFontData()
		
	fontData['FontFileName'] = os.path.basename(path)

	#Get data from font NAME table
	try:
		if len(font.name.getNameFromID(4,"")) > 1:
			fontData['Name'] = remove_control_characters(font.name.getNameFromID(4,""))
		else:
			fontData['Name'] = os.path.splitext(fileName)[0]
				
		fontData['Copyright'] = remove_control_characters(font.name.getNameFromID(0,""))
		fontData['Trademark'] = remove_control_characters(font.name.getNameFromID(7,""))
		fontData['FileType'] = os.path.splitext(path)[-1]
		fontData['Path'] = path 
		
		fontData['Family'] = remove_control_characters(font.name.getNameFromID(1,""))
		fontData['Subfamily'] = remove_control_characters(font.name.getNameFromID(2,""))
		fontData['LicenseDesc'] = remove_control_characters(font.name.getNameFromID(13,""))
		fontData['LicenseURL'] = remove_control_characters(font.name.getNameFromID(14,""))
		fontData['UniqueID'] = remove_control_characters(font.name.getNameFromID(3,""))
		fontData['Misc1'] = remove_control_characters(font.name.getNameFromID(200,""))
		fontData['Misc2'] = remove_control_characters(font.name.getNameFromID(55555,""))
		fontData['Version'] = remove_control_characters(font.name.getNameFromID(5,""))
		fontData['Manufacturer'] = remove_control_characters(font.name.getNameFromID(8,""))
		fontData['Designer'] = remove_control_characters(font.name.getNameFromID(9,""))
		fontData['DesignerURL'] = remove_control_characters(font.name.getNameFromID(12,""))
		fontData['Vendor'] = remove_control_characters(font.name.getNameFromID(11,""))
		fontData['VendorURL'] = remove_control_characters(font.name.getNameFromID(11,""))
		fontData['VendorID'] = font['OS/2'].achVendID
		
		#path is either a path or '' (for the case of single fonts)
		if path != '':
			fontData['UseCase'] = getUseCase(path)##################
		
			if 'Mobile App' in fontData['UseCase']:
				if path in arvrFonts:
					fontData['UseCase'] = 'AR/VR ' + fontData['UseCase']

				if 'Payload' in path:
					fontData['AppPlatform'] = 'iOS'
					plistpath = path
					while 'Payload' in plistpath:
						plistpath = os.path.dirname(plistpath)
					plistpath = os.path.join(plistpath,'iTunesMetadata.plist')
					try:
						with open(plistpath,'rb') as fp:
							pl = plistlib.load(fp)
							fontData['Title'] = pl["itemName"]
							fontData['AppDev'] = pl["playlistName"]
							fontData['AppURL'] = "https://apps.apple.com/app/id"+str(pl["itemId"])
					except: #if plistfile is empty...
						found = False
						p = Path(path)
						for i in range(len(p.parts)):
							if found:
								break
							if 'app_research' in p.parts[i]:
								found = True
								fontData['Title'] = p.parts[i+1]
								fontData['AppDev'] = ''
								fontData['AppURL'] = ''
				else:
					fontData['AppPlatform'] = 'Android'
					found = False
					p = Path(path)
					isAmazonApp = False
					for i in range(len(p.parts)):
						if found:
							break
						if 'app_research' in p.parts[i] or 'amazon_research' in p.parts[i]:
							if 'amazon_research' in p.parts[i]:
								isAmazonApp = True
							found = True
							fontData['Title'] = p.parts[i+1]
					if isAmazonApp:
						namepsace = fontData['Title'].split('_')[-1]
						fontData['AppURL'] = 'http://www.amazon.com/gp/mas/dl/android?p='+namepsace
					else:
						fontData['AppURL'] = 'https://play.google.com/store/apps/details?id='+fontData['Title']
					
			elif fontData['UseCase'] == 'Web Font' or fontData['UseCase'] == 'Digital Ad':
				found = False
				p = Path(path)
				for i in range(len(p.parts)):
					if found:
						break
					if 'web_research' in p.parts[i] or 'ad_research' in p.parts[i]:
						found = True
						fontData['Title'] = p.parts[i+1]
		
	except Exception as e:
	#print ("Error in ", path, fontData['FontFileName'])
		pass

	#returns boolean corresponding to likelihood of being a Monotype font
	checkFontFoundry(fontData)
	
	#GetRemarks from font metadata
	fontData['Remarks'] = getRemarks(fontData)
	
	return fontData

def extract_pdf(folder, file, logFile, mupdf_path):
	newpath = os.path.join(folder, os.path.splitext(file)[0].strip())
	if os.path.exists(newpath):
		logFile.write("pdf already extracted for:" + newpath + '.pdf' + "\n")
		return mupdf_path
	logFile.write("Extracting PDF:" + newpath + '.pdf' + "\n")
	if mupdf_path == '':
		mupdf_path = pyxj.askOpenFile(title="Select the mutool.exe:", filter=".exe")
		logFile.write("mupdf_path:" + mupdf_path + "\n")
	if not mupdf_path.endswith('mutool.exe'):
		logFile.write("invalid mupdf_path:" + mupdf_path + "\n")
		logFile.write("pdf not extracted for:" + os.path.join(folder, file) + "\n")
		mupdf_path = ''
		return mupdf_path
	cwd = os.getcwd()
	try:
		os.mkdir(newpath)
	except Exception:
		pass
	os.chdir(newpath)
	shutil.copyfile(os.path.join(folder, file), os.path.join(newpath, file))
	cmd = mupdf_path + ' extract "' + file + '"'
	logFile.write("extract Command:" + cmd + "\n")
	os.system(cmd)
	os.remove(os.path.join(newpath, file))
	os.chdir(cwd)
	return mupdf_path

def unzip(path, logFile, assetReader):
	mupdf_path = ''
	for root, dirs, files in os.walk(path):
		files = assetReader.mergeSplitFiles(root, files)
		for file in files:
			filepath = os.path.join(root, file)
			fileext = os.path.splitext(file)[-1]
			if fileext.lower() in [".apk", ".ipa", ".zip"] and os.path.splitext(file)[0] not in dirs:
				logFile.write("Extracting zip file:" + filepath + "\n")
				currentdir = os.path.join(root, os.path.splitext(file)[0])
				try:
					with ZipFile(filepath, 'r') as zipObj:
						#print("now extract\n")
						zipObj.extractall(path=currentdir)
						#print("now unzip subfolder\n")
						unzip(currentdir, logFile, assetReader)
					#print("now remove zip\n")
					#os.remove(filepath)
					logFile.write("Removed zipfile " + filepath + "\n")
				except:
					logFile.write("Exception zipfile " + filepath + "\n")
			elif '.assets' in fileext.lower() or fileext.lower() == '':
				logFile.write("Searching font in:" + filepath + "\n")
				with open(filepath, 'rb', 0) as file:
					assetReader.currentFolder = root
					if assetReader.parseSingleFile(file):
						logFile.write("Font is present in:" + filepath + "\n")
			elif '.pdf' in fileext.lower() and 'pdf_research' in filepath:
				mupdf_path = extract_pdf(root, file, logFile, mupdf_path)

def normalizeString(text):
	return re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', text)

#input: directory to scan for fonts
#output: list of fonts as dictionaries
#Note: Need to add ability to:
#			load .dfont fonts as Editors
#			extract fonts from .asset files
def scan_dir(dirpath):
	logFile = os.path.join(dirpath,'scan.log')
	logF = open(logFile, "w", encoding="utf-8")
	logF.write("Log File Created\n")

	# Unzip any zipped file
	arvrFonts = []
	reader = UnityFontAssetReader(arvrFonts)
	unzip(dirpath,logF, reader)

	if os.path.exists(os.path.join(dirpath, 'temp')):
		shutil.rmtree(os.path.join(dirpath, 'temp'))
	os.makedirs(os.path.join(dirpath, 'temp'))
	temp_path = os.path.join(dirpath, 'temp')

	fonts_list = []
	
	#list of strings to hold error messages so it doesn't spam the console for each font in the ttc
	ttc_errors = []
	
	#Holds font types not currently extractable
	#['Type',path]
	#Type: e.g. 'AR/VR Font', 'DFONT Collection'
	other_fonts = []
	
	#exclude created temp file path from scan
	exclude_dirs = ['temp']

	count = 1
	rootDirs = []

	for root, dirs, files in os.walk(dirpath, topdown = True):
		dirs[:] = [d for d in dirs if d not in exclude_dirs]

		if os.path.samefile(os.path.dirname(root), dirpath):
			rootDirs = [os.path.basename(d) for d in dirs]

		for file in files:
			fileext = os.path.splitext(file)[-1]
			filepath = os.path.join(root, file)
			temp_filepath = os.path.join(temp_path, 'file' + str(count))
			#print('Temp file path:', temp_filepath)
			#shutil.copy(filepath, temp_filepath)

			count = count + 1
			try:
				os.remove(temp_filepath)
			except OSError:
				pass

			rootDir = ""
			fontFound = False
			dirsInPath = os.path.relpath(filepath, start=dirpath).split(os.path.sep)
			if len(dirsInPath) > 1:
				rootDir = dirsInPath[1]

			if fileext.lower() in ('.ttf','.otf'):
				try:
					if file not in os.listdir(temp_path):
						shutil.copy(filepath, temp_filepath)
						#print('List dir',os.listdir(temp_path))
					editor = fontedit.Editor.frompath(temp_filepath)
					dict_font = getFontData(editor, filepath, arvrFonts,file)
					fonts_list.append(dict_font)
					logF.write("Font File Found:" + filepath + "\n")
					fontFound = True
				except Exception as err:
					logF.write("Exception reading font:" + filepath + "\n")
					pass
			
			elif fileext.lower() in ('.eot', '.woff', '.woff2'):
				try:
					if file not in os.listdir(temp_path):
						shutil.copy(filepath, temp_filepath)
						#print('List dir',os.listdir(temp_path))
					editor = fontEditorFromFont(temp_filepath, retainAllTables=True)
					dict_font = getFontData(editor, filepath, arvrFonts,file)
					fonts_list.append(dict_font)
					fontFound = True
					logF.write("Font File Found:" + filepath + "\n")
				except Exception as e:
					print('Exception', e)
					pass

			elif fileext.lower() in ('.cff') and fileext.lower() != '':
				try:
					fontData = initFontData()
					fw = FileWalker(filepath)
					cff = CFF.CFF.fromvalidatedwalker(fw)
					if cff:
						if 'notice' in cff.fontinfo:
							fontData['Copyright'] = cff.fontinfo['notice']
							if isinstance(fontData['Copyright'], (bytes, bytearray)):
								fontData['Copyright'] = fontData['Copyright'].decode("ascii", 'ignore')
							fontData['Copyright'] = normalizeString(fontData['Copyright'])
						if 'weight' in cff.fontinfo:
							fontData['Subfamily'] = cff.fontinfo['weight']
							if isinstance(fontData['Subfamily'], (bytes, bytearray)):
								fontData['Subfamily'] = fontData['Subfamily'].decode("ascii", 'ignore')
							fontData['Subfamily'] = normalizeString(fontData['Subfamily'])
						if 'uniqueID' in cff.fontinfo:
							fontData['UniqueID'] = cff.fontinfo['uniqueID']
							if isinstance(fontData['UniqueID'], (bytes, bytearray)):
								fontData['UniqueID'] = fontData['UniqueID'].decode("ascii", 'ignore')
							else:
								fontData['UniqueID'] = str(fontData['UniqueID'])
							fontData['UniqueID'] = normalizeString(fontData['UniqueID'])
						fontData['Name'] = cff.fontinfo.fontname
						if isinstance(fontData['Name'], (bytes, bytearray)):
							fontData['Name'] = fontData['Name'].decode("ascii", 'ignore')
						fontData['Name'] = normalizeString(fontData['Name'])
						fontData['Family'] = fontData['Name']
						fontData['UseCase'] = getUseCase(filepath)
						fontData['Path'] = filepath
						fontData['FontFileName'] = os.path.basename(filepath)
						checkFontFoundry(fontData)
						fontFound = True
						logF.write("Font File Found:" + filepath + "\n")
						fonts_list.append(fontData)
				except TypeError:
					print("Exception in CFF:" + filepath)

			elif fileext.lower() in ('.ttc') and fileext.lower() != '':
				try:
					if file not in os.listdir(temp_path):
						shutil.copy(filepath, temp_filepath)
						#print('List dir',os.listdir(temp_path))
					myEditorList = collectionedit.CollectionEditor.frompath(temp_filepath)
					for editor in myEditorList:
						dict_font = getFontData(editor, filepath, arvrFonts,file)
						fonts_list.append(dict_font)
					fontFound = True
					logF.write("Font File Found:" + filepath + "\n")
				except Exception:
					error_msg = 'Error unpacking ' + filepath + ' TrueType Collection. Please manually unpack using PyFontChef TTC tools and re-run this script.'
					if error_msg not in ttc_errors:
						print('Error unpacking ' + filepath + ' TrueType Collection. Please manually unpack using PyFontChef TTC tools and re-run this script.')
						ttc_errors.append(error_msg)
					pass
					
			#elif '.assets' in fileext.lower() or '.split' in fileext.lower() and fileext.lower()!='':
				#logF.write(".assets or .split is present in " + rootDir + "\n")
				#other_fonts.append(['AR/VR Font', filepath])
				#with open(filepath, 'rb', 0) as file, \
					#mmap.mmap(file.fileno(), 0, access=mmap.ACCESS_READ) as s:
					#if s.find(b'copyright') != -1:
					#other_fonts.append(['AR/VR Font', filepath])
			
			elif ('.dfont') in fileext.lower() and fileext.lower()!='':
				logF.write(".dfont is present in " + rootDir + "\n")
				other_fonts.append(['DFONT Collection', filepath])
			elif ('.fnt') in fileext.lower() and fileext.lower()!='':
				logF.write(".fnt is present in " + rootDir + "\n")
				other_fonts.append(['Bitmap Font file', filepath])
			elif ('.mvec') in fileext.lower() and fileext.lower()!='':
				logF.write(".mvec is present in " + rootDir + "\n")
				other_fonts.append(['Playground SDK Font file', filepath])

			if fontFound:
				if rootDir in rootDirs:
					rootDirs.remove(rootDir)
					logF.write(rootDir + " has font file\n")

	#Print Log for font file missing folders
	logF.write("\n\n\n")
	for rootDir in rootDirs:
		print("Font is not present in:" + rootDir)
		logF.write("Font is not present in:" + rootDir + "\n")

	#Convert fonts & paths to dicts
	#store font dicts in fonts_list
	
	for other_font in other_fonts:
		fontData = {}
		fontData['UseCase'] = getUseCase(other_font[1])
				
		if 'Mobile App' in fontData['UseCase']:
			if 'Payload' in other_font[1]:
				fontData['AppPlatform'] = 'iOS'
				plistpath = other_font[1]
				while 'Payload' in plistpath:
					plistpath = os.path.dirname(plistpath)
				plistpath = os.path.join(plistpath,'iTunesMetadata.plist')
				try:
					with open(plistpath,'rb') as fp:
						pl = plistlib.load(fp)
						fontData['Title'] = pl["itemName"]
						fontData['AppDev'] = pl["playlistName"]
						fontData['AppURL'] = "https://apps.apple.com/app/id"+str(pl["itemId"])
				except: #if plistfile is empty...
					found = False
					p = Path(other_font[1])
					for i in range(len(p.parts)):
						if found:
							break
						if 'app_research' in p.parts[i]:
							found = True
							fontData['Title'] = p.parts[i+1]
			else:
				fontData['AppPlatform'] = 'Android'
				found = False
				p = Path(other_font[1])
				for i in range(len(p.parts)):
					if found:
						break
					if 'app_research' in p.parts[i]:
						found = True
						fontData['Title'] = p.parts[i+1]
				fontData['AppURL'] = 'https://play.google.com/store/apps/details?id='+fontData['Title']
		
		fontData['Name'] = other_font[0]
		fontData['Path'] = other_font[1]
	
		fonts_list.append(fontData)

	logF.close()
	return fonts_list

def create_output_folder(dir):
	newName = 'documentation'
	allDirs = [o for o in os.listdir(dir)
                    if os.path.isdir(os.path.join(dir, o)) and o.startswith(newName)]
	if len(allDirs) > 0:
		allDirs.sort()
		if len(allDirs) == 1 and allDirs[0] == newName:
			index = 0
		else:
			index = allDirs[len(allDirs) -1][len(newName):]
		try:
			newName = newName + str(int(index) + 1)
		except:
			pass
	fullPath = os.path.join(dir, newName)
	if os.path.exists(fullPath):
		shutil.rmtree(fullPath)
	os.makedirs(fullPath)
	return fullPath

def _main():
	#Prompt what the script should do - scan, review website scrape results, or create documentation
	my_dialog = dialog.Dialog(title = 'Font Use Wizard ' + SCRIPT_VERSION, cancellable = True)
	label_top = dialog.Label('What would you like to do?')
	label_scan = dialog.Label('Scan Tools: Scan a directory for fonts or review a font for details')
	label_scraper = dialog.Label('Web Scraper Tools: Review/validate Gutenberg web scrape results')
	label_document = dialog.Label('Report Tools: Create an Account Plan, Font Use CSV, or Font Use Doc')
	label_apps = dialog.Label('App Tools: Check which countries iOS apps are available in')
	my_dialog.add(label_top)
	my_dialog.add(dialog.Separator(line=True))
	rbg = dialog.RadiobuttonGroup(['Scan Tools', 'Scraper Tools','App Tools','Report Tools'],selectedindex = 0)
	my_dialog.add(rbg)
	my_dialog.add(label_scan)
	my_dialog.add(label_scraper)
	my_dialog.add(label_apps)
	my_dialog.add(label_document)
	

	#script not cancelled
	if dialog.show(my_dialog):
		#scan
		if rbg.selectedindex == 0:
			scan_dialog = dialog.Dialog(title = 'Font Scanner ' + SCRIPT_VERSION, cancellable = True)
			label_top = dialog.Label('Please choose an option')
			label_opt1 = dialog.Label('Scan Directory: Scans a directory for fonts and creates a report')
			#label_opt2 = dialog.Label('Scan Font: Select a Font')
			#label_opt3 = dialog.Label('Scan Font: Loaded Fonts')
			scan_dialog.add(label_top)
			scan_dialog.add(dialog.Separator(line=True))
			rbg = dialog.RadiobuttonGroup(['Scan Directory'],selectedindex = 0)
			#rbg = dialog.RadiobuttonGroup(['Scan Directory', 'Scan Selected Font', 'Scan Loaded Fonts'],selectedindex = 0)
			scan_dialog.add(rbg)
			scan_dialog.add(label_opt1)
			#scan_dialog.add(label_opt2)
			#scan_dialog.add(label_opt3)
			
			if dialog.show(scan_dialog):
				#Scan Directory
				if rbg.selectedindex == 0:
					basic_dialog = dialog.Dialog(title = 'Select Scan Data ' + SCRIPT_VERSION)
					label_scan = dialog.Label('Choose a directory to scan for fonts')
					basic_dialog.add(label_scan)
					
					if dialog.show(basic_dialog):
						dirpath = pyxj.askOpenFolder(title = 'Choose a directory to scan for fonts.')
					fonts_list = scan_dir(dirpath)
					
					
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_save = dialog.Label('Choose would you like to save the scan results')
					basic_dialog.add(label_save)
					
					if dialog.show(basic_dialog):
						outpath = pyxj.askOpenFolder(title = 'Choose would you like to save the scan results?')
						
					write_dir_scan(fonts_list, outpath)
				elif rbg.selectedindex == 1:
					try:
						font_path = pyxj.askOpenFile(title="Select a font")
						editor = fontedit.Editor.frompath(font_path)
						fontData = getFontData(editor,font_path, [],font_path)
						
						print(len(fontData.keys))
						
						for key in fontData:
							print('key:' + key)
							print('value: ' + key.value)
						
					except Exception:
						print('Could not read font!')
				elif rbg.selectedindex == 2:
					#will cycle through loaded Editors
					pass
					
				
				#dirpath = pyxj.askOpenFolder(title = 'Choose a directory to scan for fonts.')
			
				#if os.path.isdir(dirpath + 'web_research') or os.path.isdir(dirpath + 'app_research):
				#	report_scan(dirpath)
				#else:
				#	gen_scan(dirpath)
		#scraper
		elif rbg.selectedindex == 1:
			scrape_dialog = dialog.Dialog(title = 'Web Scraper ' + SCRIPT_VERSION, cancellable = True)
			label_top = dialog.Label('Please choose an option')
			label_opt1 = dialog.Label('Full Validation of Scrape Results- check for errors, missed websites, and redirects')
			label_opt2 = dialog.Label('Check redirects e.g. input: mywebsite.com, scrape: facebook.com/mywebsite')
			label_opt3 = dialog.Label('Check Missed Websites: Did the scraper miss any input websites not due to error?')
			label_opt4 = dialog.Label('Check Web Scraper Errors: Get a list of errors in the scrape')
			scrape_dialog.add(label_top)
			scrape_dialog.add(dialog.Separator(line=True))
			rbg = dialog.RadiobuttonGroup(['Full Validation','Check Redirects', 'Check Missed Websites', 'Check Web Scraper Errors'],selectedindex = 0)
			scrape_dialog.add(rbg)
			scrape_dialog.add(label_opt1)
			scrape_dialog.add(label_opt2)
			scrape_dialog.add(label_opt3)
			scrape_dialog.add(label_opt4)
			
			if dialog.show(scrape_dialog):
				if rbg.selectedindex == 0:
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_prescrape = dialog.Label('Select the pre-scrape website CSV')
					basic_dialog.add(label_prescrape)
					
					#prompt for input website csv and gutenberg results csv
					if dialog.show(basic_dialog):
						input_websites_path = pyxj.askOpenFile(title="Select the pre-scrape website CSV:", filter=".csv")
					
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_scrape = dialog.Label('Select the Web Scraper Results CSV')
					basic_dialog.add(label_scrape)
					
					if dialog.show(basic_dialog):
						scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
					
					#check for errors in web scrape
					scrape_errors = getScrapeErrors(scrape_results_path)
					
					#check for missed websites by comparing input list and scrape
					missed_websites = checkMissedWebsites(input_websites_path, scrape_results_path)
					
					#output list of errors in web scrape
					if not scrape_errors:
						print('Error Check: The web scraper did not encounter any errors.\n')
					else:
						print('Error Check: The web scraper encountered errors for the following websites:')
						for scrape_error in scrape_errors:
							print('Domain: ' + scrape_error[0] + '\nError Type: ' + scrape_error[1] + '\n')
						print('The above websites require manual review.\n')
					
					#output list of missed websites
					if not missed_websites:
						print('Missed Website Check: No websites missed.\n')
					else:
						print('Missed Website Check: The following websites were not read in by the web scraper and may need to be manually reviewed:\n')
						for missed_website in missed_websites:
							print(missed_website)
										
					#check for redirects and confirm whether to include in report
					#this catches issues where the scraper draws in some unrelated link e.g. facebook when crawling
					redirect_websites = getRedirects(scrape_results_path)
					
				elif rbg.selectedindex == 1:
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_scrape = dialog.Label('Select the Web Scraper Results CSV')
					basic_dialog.add(label_scrape)
					
					if dialog.show(basic_dialog):
						scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
					
					redirect_websites = getRedirects(scrape_results_path)
					
				elif rbg.selectedindex == 2:
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_prescrape = dialog.Label('Select the pre-scrape website CSV')
					basic_dialog.add(label_prescrape)
					
					if dialog.show(basic_dialog):
						input_websites_path = pyxj.askOpenFile(title="Select the pre-scrape website CSV:", filter=".csv")
					
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_scrape = dialog.Label('Select the Web Scraper Results CSV')
					basic_dialog.add(label_scrape)
					
					if dialog.show(basic_dialog):
						scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
					
					missed_websites = checkMissedWebsites(input_websites_path, scrape_results_path)
					
					if not missed_websites:
						print('Missed Website Check: No websites missed.\n')
					else:
						print('Missed Website Check: The following websites were not read in by the web scraper and may need to be manually reviewed:\n')
						for missed_website in missed_websites:
							print(missed_website)
							
				elif rbg.selectedindex == 3:
					basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
					label_scrape = dialog.Label('Select the Web Scraper Results CSV')
					basic_dialog.add(label_scrape)
					
					if dialog.show(basic_dialog):
						scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
					scrape_errors = getScrapeErrors(scrape_results_path)
					
					if not scrape_errors:
						print('Error Check: The web scraper did not encounter any errors.\n')
					else:
						print('Error Check: The web scraper encountered errors for the following websites:')
						for scrape_error in scrape_errors:
							print('Domain: ' + scrape_error[0] + '\nError Type: ' + scrape_error[1] + '\n')
						print('The above websites require manual review.\n')
						
		#app tools
		elif rbg.selectedindex == 2:
			app_dialog = dialog.Dialog(title = 'App Tools ' + SCRIPT_VERSION, cancellable = True)
			label_top = dialog.Label('Please choose an option')
			label_opt1 = dialog.Label('iOS Country Finder: Generates a CSV file with a list of countries where each app can be downloaded')
			app_dialog.add(label_top)
			rbg = dialog.RadiobuttonGroup(['iOS Country Finder'],selectedindex = 0)
			app_dialog.add(rbg)
			app_dialog.add(label_opt1)
			
			if dialog.show(app_dialog):
				if rbg.selectedindex == 0:
					#inputs = dev_url, save_location, current_country
					#developer url
					dev_url_dialog = dialog.Dialog(title = 'Developer URL ' + SCRIPT_VERSION, cancellable = False)
					label_dev1 = dialog.Label("Enter the iOS developer URL. Here's an example: https://apps.apple.com/developer/pandora-media-inc/id284035180")
					label_dev2 = dialog.Label("Note: format must match exactly. A common mistake would be e.g. https://apps.apple.com/developer/pandora-media-inc/id284035180 \nbecause it includes /us/")
					dev_ef = dialog.Editfield(initial = "https://apps.apple.com/developer/[DeveloperName]/[DeveloperID]")
					
					dev_url_dialog.add(label_dev1)
					dev_url_dialog.add(label_dev2)
					dev_url_dialog.add(dev_ef)
					
					if dialog.show(dev_url_dialog):
						dev_url = dev_ef.text
					
					#current_country
					cc_dialog = dialog.Dialog(title = 'Current Country ' + SCRIPT_VERSION, cancellable = False)
					label_cc1 = dialog.Label("Enter your current country.\nThe results will show whether an app is available in your local store or elsewhere.\nIf you are connected through a VPN to USA, enter United States")
					cc_ef = dialog.Editfield(initial = "United States")
					
					cc_dialog.add(label_cc1)
					cc_dialog.add(cc_ef)
					
					if dialog.show(cc_dialog):
						cc = cc_ef.text
					
					save_dialog = dialog.Dialog(title = 'Save Location ' + SCRIPT_VERSION, cancellable = False)
					label_save = dialog.Label("Choose a location to save the results. \nThey will be saved in that directory to a file called iOS_countries.csv")
					save_dialog.add(label_save)
					
					if dialog.show(save_dialog):
						output_folder_path = pyxj.askOpenFolder(title = 'Choose where would you like to save the results')

					find_iOS_apps(dev_url,cc,output_folder_path)
					
			else:
				print('Script Cancelled!')
		#documentation	
		elif rbg.selectedindex == 3:
			doc_dialog = dialog.Dialog(title = 'Documentation ' + SCRIPT_VERSION, cancellable = True)
			label_top = dialog.Label('Please choose an option')
			label_opt1 = dialog.Label('Full Account Plan: Account Plan CSVs, Font Use CSV, & Font Use Doc')
			label_opt2 = dialog.Label('Font Use CSV: CSV for use in customizing Font Use Docs')
			label_opt3 = dialog.Label('Font Use Doc: Create documentation based on Font Use CSV')
			doc_dialog.add(label_top)
			doc_dialog.add(dialog.Separator(line=True))
			rbg = dialog.RadiobuttonGroup(['Full Account Plan', 'Font Use CSV', 'Font Use Doc'],selectedindex = 0)
			doc_dialog.add(rbg)
			doc_dialog.add(label_opt1)
			doc_dialog.add(label_opt2)
			doc_dialog.add(label_opt3)
		
			if dialog.show(doc_dialog):
				if rbg.selectedindex == 0:
					#prompt for update, new and whether there is a scrape csv
					cb_dialog = dialog.Dialog(title = 'Account Planning ' + SCRIPT_VERSION, cancellable = True)
					cb = dialog.Checkbox('I have web scrape results I want to include', value=True)
					cb_dialog.add(cb)
					pre_csv_path = ''
					if dialog.show(cb_dialog):
						
						basic_dialog = dialog.Dialog(title =  SCRIPT_VERSION)
						label_scan = dialog.Label('Choose a directory to scan for fonts.')
						basic_dialog.add(label_scan)
					
						if dialog.show(basic_dialog):
							dirpath = pyxj.askOpenFolder(title = 'Choose a directory to scan for fonts.')
						
						fonts_list = scan_dir(dirpath)

						redirectedURLs = {}
						if cb.value:
							basic_dialog = dialog.Dialog(title =  SCRIPT_VERSION)
							label_scrape = dialog.Label('Select the Web Scraper Results CSV')
							basic_dialog.add(label_scrape)
					
							if dialog.show(basic_dialog):
								scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
							#parsed_scrape = [fonts_list, websites_not_using_webfonts]
							parsed_scrape = readScrapeResults(scrape_results_path)
							redirectedURLs = parsed_scrape[2]
							fonts_list = fonts_list + parsed_scrape[0]

							cb_dialog = dialog.Dialog(title='Account Planning ' + SCRIPT_VERSION, cancellable=True)
							cb = dialog.Checkbox('I want to use pre-generated FontInfo.csv to create final report', value=True)
							cb_dialog.add(cb)
							if dialog.show(cb_dialog) and cb.value:
								basic_dialog = dialog.Dialog(title =  SCRIPT_VERSION)
								label_scrape = dialog.Label('Select the pre-generated FontInfo.csv')
								basic_dialog.add(label_scrape)
								if dialog.show(basic_dialog):
									pre_csv_path = pyxj.askOpenFile(title="Select the pre-generated FontInfo.csv:",
																		   filter=".csv")

						output_folder_path = os.path.join(create_output_folder(dirpath), '.')
						
						#create font use csv, save it to scanned directory
						csv_path = output_folder_path + 'FontInfo_' + SCRIPT_VERSION + '.csv'
						create_font_use_csv(fonts_list, csv_path)
					
						#create font use doc from font use csv
						create_font_use_doc(csv_path, output_folder_path, pre_csv_path, redirectedURLs)
						
						#create Account Plan website tab
						#if cb.value:
						#	create_ap_web(csv_path, output_folder_path, parsed_scrape[1])
						#else:
						#	create_ap_web(csv_path, output_folder_path)
						create_ap_web(csv_path, output_folder_path)
						
						#create Account Plan apps tab
						create_ap_apps(csv_path, output_folder_path)
						
						print('Done!')
						print('Documentation can be found at ' + output_folder_path)
						
					else:
						print('Script Cancelled!')
						
				elif rbg.selectedindex == 1:
					cb_dialog = dialog.Dialog(title = 'Account Planning ' + SCRIPT_VERSION, cancellable = True)
					cb = dialog.Checkbox('I have web scrape results I want to include', value=True)
					cb_dialog.add(cb)
					scrape_fonts_list = []
					if dialog.show(cb_dialog):
						basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
						label_scan = dialog.Label('Choose a directory to scan for fonts.')
						basic_dialog.add(label_scan)
					
						if dialog.show(basic_dialog):
							dirpath = pyxj.askOpenFolder(title = 'Choose a directory to scan for fonts.')
						fonts_list = scan_dir(dirpath)
						
						if cb.value:
							basic_dialog = dialog.Dialog(title = SCRIPT_VERSION)
							label_scrape = dialog.Label('Select the Web Scraper Results CSV')
							basic_dialog.add(label_scrape)
					
							if dialog.show(basic_dialog):
								scrape_results_path = pyxj.askOpenFile(title="Select the Web Scraper Results CSV:", filter=".csv")
							#parsed_scrape = [fonts_list, websites_not_using_webfonts]
							parsed_scrape = readScrapeResults(scrape_results_path)
							fonts_list = fonts_list + parsed_scrape[0]
						
						output_folder_path = create_output_folder(dirpath) + '\\'

						#create font use csv, save it to scanned directory
						csv_path = output_folder_path + 'FontInfo_' + SCRIPT_VERSION + '.csv'
						create_font_use_csv(fonts_list,csv_path)
					
						print('Done!')
						print('Documentation can be found at ' + output_folder_path)
					
					else:
						print('Script Cancelled!')
						
				elif rbg.selectedindex == 2:
					basic_dialog = dialog.Dialog(title =  SCRIPT_VERSION)
					label_select = dialog.Label('Select a Font Use CSV')
					basic_dialog.add(label_select)
					
					if dialog.show(basic_dialog):
						csv_path = pyxj.askOpenFile(title="Select a Font Use CSV:", filter=".csv")
						
					basic_dialog = dialog.Dialog(title =  SCRIPT_VERSION)
					label_save = dialog.Label('Choose where would you like to save the scan results')
					basic_dialog.add(label_save)
					
					if dialog.show(basic_dialog):	
						output_folder_path = pyxj.askOpenFolder(title = 'Choose where would you like to save the scan results?')
					create_font_use_doc(csv_path, output_folder_path)
					
					print('Done!')
					print('Documentation can be found at ' + output_folder_path)

	else:
		print('Script cancelled!')


class UnityFontAssetReader:
	# This code is written following the logic implemented in https://github.com/Perfare/AssetStudio
    def __init__(self, fontList, folder = ""):
        self.root = folder
        self.fontList = fontList

    def mergeSplitFiles(self, folder, files):
        splitfiles = [file for file in files if os.path.splitext(file)[-1] == '.split0']
        for file in splitfiles:
            splitno = 0
            filename = os.path.splitext(file)[0]
            openfile = open(os.path.join(folder, filename), "wb")
            while filename + '.split' + str(splitno) in files:
                splitfile = open(os.path.join(folder, filename + '.split' + str(splitno)), "rb")
                openfile.write(splitfile.read())
                splitfile.close()
                splitno = splitno + 1
            openfile.close()
            files.append(filename)
        return files

    def readTypeTree(self, file, version, level=0):
        type = ''.join(iter(lambda: file.read(1).decode('ascii'), '\x00'))
        name = ''.join(iter(lambda: file.read(1).decode('ascii'), '\x00'))
        file.seek(12, 1)
        if version == 2:
            file.seek(4, 1)
        if version != 3:
            file.seek(8, 1)
        count, = struct.unpack(format + "i", file.read(4))
        for j in range(count):
            self.readTypeTree(file, version, level + 1)

    def align(self, file):
        pos = file.tell()
        mod = pos % 4
        if mod != 0:
            file.seek(4 - mod, 1)

    def getFontData(self, file, format, unityVersion, headerVersion):
        fontNameLength, = struct.unpack(format + "i", file.read(4))
        fontName = []
        pos = file.tell()
        file.seek(0, 2)
        size = file.tell()
        file.seek(pos)
        if fontNameLength > 0 and fontNameLength <= size - pos:
            fontName = file.read(fontNameLength).decode()
            self.align(file)

        version = re.sub(r"\D", ".", unityVersion).split('.')
        version = [int(i) for i in version]
        fontData = []
        if (version[0] == 5 and version[1] >= 5) or version[0] > 5:
            if headerVersion < 14:
                file.seek(44,1)
            else:
                file.seek(52,1)
            characterRects_size, = struct.unpack(format + "i", file.read(4))
            for i in range(characterRects_size):
                file.seek(44, 1)
            kerningValues_size, = struct.unpack(format + "i", file.read(4))
            for i in range(kerningValues_size):
                file.seek(8, 1)
            file.seek(4,1)
            fontData_size, = struct.unpack(format + "i", file.read(4))
            if fontData_size > 0:
                fontData = file.read(fontData_size)
        else:
            file.seek(12,1)
            if version[0] <= 3:
                file.seek(8,1)
                size, = struct.unpack(format + "i", file.read(4))
                for i in range(size):
                    file.seek(8,1)
            else:
                file.seek(8,1)
            if headerVersion < 14:
                file.seek(12,1)
            else:
                file.seek(16,1)
            size, = struct.unpack(format + "i", file.read(4))
            for i in range(size):
                file.seek(40,1)
                if version[0] >= 4:
                    file.seek(1,1)
                    self.align(file)
            if headerVersion < 14:
                file.seek(8,1)
            else:
                file.seek(12,1)
            size, = struct.unpack(format + "i", file.read(4))
            for i in range(size):
                file.seek(8,1)
            if version[0] <= 3:
                file.seek(1,1)
                self.align(file)
            else:
                file.seek(4,1)
            fontData_size, = struct.unpack(format + "i", file.read(4))
            if fontData_size > 0:
                fontData = file.read(fontData_size)
        return fontName, fontData

    def saveFont(self, file, format, unityVersion, headerVersion):
        fontName, fontData = self.getFontData(file, format, unityVersion, headerVersion)
        if fontName != "" and len(fontData) > 0:
            fontPath = os.path.join(self.currentFolder, fontName + ".ttf")
            self.fontList.append(fontPath)
            with open(fontPath, "wb") as fontFile:
                fontFile.write(fontData)

    def parseSingleFile(self, file):
        isFontPresent = False
        try:
            file.seek(0)
            metadataSize, = struct.unpack(">I", file.read(4))
            fileSize, = struct.unpack(">I", file.read(4))
            version, = struct.unpack(">I", file.read(4))
            offset, = struct.unpack(">I", file.read(4))
            if version > 50 or fileSize < metadataSize or fileSize < offset:
                return False

            if version >= 9:
                endian = int.from_bytes(file.read(1), byteorder='big')
                file.seek(3,1)
            else:
                file.seek(fileSize - metadataSize)
                endian = file.read(1)

            format = ">"
            if endian == 0:
                format = "<"
            elif endian != 1:
                return False

            if version >= 22:
                file.seek(28, 1)

            if version >= 7:
                unityVersion = ''.join(iter(lambda: file.read(1).decode('ascii'), '\x00'))

            if version >= 8:
                file.seek(4, 1)

            typeTreeEnabled = False
            if version >= 13:
                typeTreeEnabled, = struct.unpack(format + "?", file.read(1))

            typeCount, = struct.unpack(format + "I", file.read(4))

            classIDs = []
            for i in range(typeCount):
                classID, = struct.unpack(format + "i", file.read(4))
                classIDs.append(classID)
                if classID == 128:
                    isFontPresent = True
                if version >= 16:
                    file.seek(1, 1)
                if version >= 17:
                    file.seek(2, 1)
                if version >= 13:
                    if (version < 16 and classID < 0) or (version >= 16 and classID == 114):
                        file.seek(16, 1)
                    file.seek(16, 1)
                if typeTreeEnabled:
                    if version >= 12 or version == 10:
                        numberOfNodes, = struct.unpack(format + "i", file.read(4))
                        stringBufferSize, = struct.unpack(format + "i", file.read(4))
                        for j in range(numberOfNodes):
                            file.seek(24, 1)
                            if version >= 19:
                                file.seek(8, 1)
                        file.read(stringBufferSize)
                    else:
                        self.readTypeTree(file, version)

            if isFontPresent:
                bigIDEnabled = 0
                if version >= 7 and version < 14:
                    bigIDEnabled, = struct.unpack(format + "i", file.read(4))

                objectCount, = struct.unpack(format + "i", file.read(4))
                for i in range(objectCount):
                    if bigIDEnabled != 0:
                        file.seek(8, 1)
                    elif version < 14:
                        file.seek(4, 1)
                    else:
                        self.align(file)
                        file.seek(8, 1)

                    if version >= 22:
                        byteStart, = struct.unpack(format + "i", file.read(8))
                    else:
                        byteStart, = struct.unpack(format + "I", file.read(4))
                    byteStart = byteStart + offset

                    file.seek(4, 1)
                    typeID,= struct.unpack(format + "i", file.read(4))
                    if version < 16:
                        classID, = struct.unpack(format + "H", file.read(2))
                    else:
                        classID = classIDs[typeID]
                    if version < 11:
                        file.seek(2, 1)
                    if version >= 11 and version < 17:
                        file.seek(2, 1)
                    if version == 15 or version == 16:
                        file.seek(1, 1)
                    if classID == 128:
                        pos = file.tell()
                        file.seek(byteStart)
                        self.saveFont(file, format, unityVersion, version)
                        file.seek(pos)
        except Exception as e:
            print(e)
        return isFontPresent

    def parse(self):
        for root, dirs, files in os.walk(self.root, topdown=False):
            for name in files:
                path = os.path.join(root, name)
                fileext = os.path.splitext(path)[-1]
                if '.assets' in fileext.lower() or '.split' in fileext.lower() or fileext.lower() == '':
                    #print("checking:" + path)
                    with open(path, 'rb', 0) as file:
                        self.currentFolder = root
                        if self.parseSingleFile(file):
                            print("Present:" + path)


if __name__ == '__main__':
    _main()
# corrected key & index errors
